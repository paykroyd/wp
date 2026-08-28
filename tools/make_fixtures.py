#!/usr/bin/env python3
"""Generate corpus fixtures. Needs python-docx for the first two; the
pathological one is built by hand so it can contain constructs python-docx
can't author (tracked changes, comments, hyperlinks, fields, content controls).
"""
import sys
import zipfile

out_dir = sys.argv[1] if len(sys.argv) > 1 else "corpus"

try:
    import docx
    from docx.enum.text import WD_BREAK
    from docx.shared import Pt, RGBColor, Inches

    d = docx.Document()
    d.add_heading("Quarterly Report", 0)
    d.add_heading("Summary", 1)
    p = d.add_paragraph("The quarterly results ")
    p.add_run("exceeded").bold = True
    p.add_run(" projections in every region except ")
    r = p.add_run("EMEA")
    r.italic = True
    r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    p.add_run(", where currency effects reduced reported growth by roughly four points.")
    d.add_heading("Regional detail", 2)
    p = d.add_paragraph()
    r = p.add_run("North America grew 14% year over year. ")
    r.font.size = Pt(14)
    r = p.add_run("Underlined")
    r.underline = True
    p.add_run(" and ")
    r = p.add_run("struck")
    r.font.strike = True
    p.add_run(" and x")
    r = p.add_run("2")
    r.font.superscript = True
    p.add_run(".\tTabbed text here.")
    d.add_paragraph("First bullet", style="List Bullet")
    d.add_paragraph("Second bullet", style="List Bullet")
    d.add_paragraph("Numbered one", style="List Number")
    d.add_paragraph("Numbered two", style="List Number")
    p = d.add_paragraph("Centered paragraph with a hanging indent.")
    p.paragraph_format.alignment = 1
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.keep_with_next = True
    t = d.add_table(rows=2, cols=3)
    t.style = "Table Grid"
    for i in range(2):
        for j in range(3):
            t.cell(i, j).text = f"r{i}c{j}"
    p = d.add_paragraph("Text after the table, before a page break.")
    p.add_run().add_break(WD_BREAK.PAGE)
    d.add_paragraph("Appendix A — Methodology")
    body = ("Lorem ipsum dolor sit amet, consectetur adipiscing elit, sed do eiusmod tempor "
            "incididunt ut labore et dolore magna aliqua. Ut enim ad minim veniam, quis nostrud "
            "exercitation ullamco laboris nisi ut aliquip ex ea commodo consequat. ")
    for i in range(40):
        d.add_paragraph(body * 2)
    d.save(f"{out_dir}/report.docx")
    print("wrote report.docx")

    d = docx.Document()
    d.save(f"{out_dir}/empty.docx")
    print("wrote empty.docx")
except ImportError:
    print("python-docx not available; skipping generated fixtures")

W = 'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml" xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006" mc:Ignorable="w14"'
DOC = f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document {W}><w:body>
<w:p w:rsidR="00A1" w14:paraId="1"><w:pPr><w:pStyle w:val="Heading1"/><w:rPr><w:lang w:val="en-GB"/></w:rPr></w:pPr><w:r><w:rPr><w:lang w:val="en-GB"/></w:rPr><w:t>Pathological</w:t></w:r></w:p>
<w:p><w:r><w:t xml:space="preserve">Plain </w:t></w:r><w:r><w:rPr><w:b/><w:bCs/><w:rFonts w:ascii="Arial" w:hAnsi="Arial" w:cs="Arial"/></w:rPr><w:t>bold Arial</w:t></w:r><w:r><w:t xml:space="preserve"> then </w:t></w:r><w:ins w:id="1" w:author="A" w:date="2024-01-01T00:00:00Z"><w:r><w:t>inserted</w:t></w:r></w:ins><w:del w:id="2" w:author="A" w:date="2024-01-01T00:00:00Z"><w:r><w:delText xml:space="preserve"> deleted</w:delText></w:r></w:del><w:r><w:t>.</w:t></w:r></w:p>
<w:p><w:commentRangeStart w:id="0"/><w:r><w:t>Commented text</w:t></w:r><w:commentRangeEnd w:id="0"/><w:r><w:commentReference w:id="0"/></w:r></w:p>
<w:p><w:r><w:t xml:space="preserve">See </w:t></w:r><w:hyperlink r:id="rId9" w:history="1"><w:r><w:rPr><w:rStyle w:val="Hyperlink"/></w:rPr><w:t>the website</w:t></w:r></w:hyperlink><w:r><w:t xml:space="preserve"> and page </w:t></w:r><w:r><w:fldChar w:fldCharType="begin"/></w:r><w:r><w:instrText xml:space="preserve"> PAGE </w:instrText></w:r><w:r><w:fldChar w:fldCharType="separate"/></w:r><w:r><w:t>3</w:t></w:r><w:r><w:fldChar w:fldCharType="end"/></w:r><w:r><w:t>.</w:t></w:r></w:p>
<w:bookmarkStart w:id="7" w:name="regional"/><w:p><w:pPr><w:jc w:val="center"/><w:ind w:left="720" w:hanging="360"/><w:spacing w:before="240" w:after="0" w:line="360" w:lineRule="auto"/><w:tabs><w:tab w:val="right" w:leader="dot" w:pos="8640"/></w:tabs></w:pPr><w:r><w:rPr><w:i/><w:sz w:val="28"/><w:color w:val="FF0000"/><w:highlight w:val="yellow"/><w:u w:val="double"/></w:rPr><w:t>Formatted</w:t></w:r><w:r><w:tab/><w:t>tabbed</w:t><w:br/><w:t>after break</w:t></w:r></w:p><w:bookmarkEnd w:id="7"/>
<w:sdt><w:sdtPr><w:alias w:val="Title"/><w:id w:val="123"/></w:sdtPr><w:sdtContent><w:p><w:r><w:t>Inside a block content control</w:t></w:r></w:p></w:sdtContent></w:sdt>
<w:p><w:r><w:t xml:space="preserve">Run-level </w:t></w:r><w:sdt><w:sdtPr><w:id w:val="124"/></w:sdtPr><w:sdtContent><w:r><w:t>control</w:t></w:r></w:sdtContent></w:sdt><w:r><w:t xml:space="preserve"> here, then a footnote</w:t></w:r><w:r><w:rPr><w:vertAlign w:val="superscript"/></w:rPr><w:footnoteReference w:id="1"/></w:r><w:r><w:t>.</w:t></w:r></w:p>
<w:tbl><w:tblPr><w:tblStyle w:val="TableGrid"/><w:tblW w:w="0" w:type="auto"/></w:tblPr><w:tblGrid><w:gridCol w:w="4000"/><w:gridCol w:w="4000"/></w:tblGrid><w:tr><w:tc><w:tcPr><w:tcW w:w="4000" w:type="dxa"/></w:tcPr><w:p><w:r><w:t>A1</w:t></w:r></w:p></w:tc><w:tc><w:tcPr><w:tcW w:w="4000" w:type="dxa"/></w:tcPr><w:p><w:r><w:t>B1</w:t></w:r></w:p></w:tc></w:tr></w:tbl>
<w:p><w:r><w:t>Last &amp; &lt;escaped&gt; paragraph</w:t></w:r><w:r><w:br w:type="page"/></w:r><w:r><w:t>New page</w:t></w:r></w:p>
<w:sectPr w:rsidR="00B2"><w:footerReference w:type="default" r:id="rId8"/><w:pgSz w:w="11906" w:h="16838"/><w:pgMar w:top="1440" w:right="1800" w:bottom="1440" w:left="1800" w:header="708" w:footer="708" w:gutter="0"/><w:cols w:space="708"/><w:docGrid w:linePitch="360"/></w:sectPr>
</w:body></w:document>'''

STYLES = f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:styles {W}><w:docDefaults><w:rPrDefault><w:rPr><w:rFonts w:asciiTheme="minorHAnsi" w:hAnsiTheme="minorHAnsi"/><w:sz w:val="24"/><w:lang w:val="en-US"/></w:rPr></w:rPrDefault><w:pPrDefault><w:pPr><w:spacing w:after="200" w:line="276" w:lineRule="auto"/></w:pPr></w:pPrDefault></w:docDefaults>
<w:latentStyles w:defLockedState="0" w:count="1"><w:lsdException w:name="Normal" w:uiPriority="0"/></w:latentStyles>
<w:style w:type="paragraph" w:default="1" w:styleId="Normal"><w:name w:val="Normal"/><w:qFormat/></w:style>
<w:style w:type="paragraph" w:styleId="Heading1"><w:name w:val="heading 1"/><w:basedOn w:val="Normal"/><w:next w:val="Normal"/><w:pPr><w:keepNext/><w:spacing w:before="480" w:after="0"/><w:outlineLvl w:val="0"/></w:pPr><w:rPr><w:rFonts w:asciiTheme="majorHAnsi" w:hAnsiTheme="majorHAnsi"/><w:b/><w:color w:val="365F91"/><w:sz w:val="28"/></w:rPr></w:style>
<w:style w:type="character" w:default="1" w:styleId="DefaultParagraphFont"><w:name w:val="Default Paragraph Font"/><w:semiHidden/></w:style>
<w:style w:type="character" w:styleId="Hyperlink"><w:name w:val="Hyperlink"/><w:rPr><w:color w:val="0000FF"/><w:u w:val="single"/></w:rPr></w:style>
</w:styles>'''

FOOTNOTES = f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:footnotes {W}><w:footnote w:type="separator" w:id="-1"><w:p><w:r><w:separator/></w:r></w:p></w:footnote><w:footnote w:id="1"><w:p><w:pPr><w:pStyle w:val="FootnoteText"/></w:pPr><w:r><w:footnoteRef/></w:r><w:r><w:t xml:space="preserve"> A footnote.</w:t></w:r></w:p></w:footnote></w:footnotes>'''

COMMENTS = f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:comments {W}><w:comment w:id="0" w:author="Reviewer" w:date="2024-01-01T00:00:00Z" w:initials="R"><w:p><w:r><w:t>Please check this.</w:t></w:r></w:p></w:comment></w:comments>'''

FOOTER = f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:ftr {W}><w:p><w:pPr><w:jc w:val="center"/></w:pPr><w:r><w:fldChar w:fldCharType="begin"/></w:r><w:r><w:instrText>PAGE</w:instrText></w:r><w:r><w:fldChar w:fldCharType="end"/></w:r></w:p></w:ftr>'''

THEME = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<a:theme xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" name="Office Theme"><a:themeElements><a:fontScheme name="Office"><a:majorFont><a:latin typeface="Cambria"/></a:majorFont><a:minorFont><a:latin typeface="Times New Roman"/></a:minorFont></a:fontScheme></a:themeElements></a:theme>'''

CT = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types"><Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/><Default Extension="xml" ContentType="application/xml"/><Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/><Override PartName="/word/styles.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.styles+xml"/><Override PartName="/word/footnotes.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml"/><Override PartName="/word/comments.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"/><Override PartName="/word/footer1.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.footer+xml"/><Override PartName="/word/theme/theme1.xml" ContentType="application/vnd.openxmlformats-officedocument.theme+xml"/></Types>'''

RELS = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"><Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/></Relationships>'''

DOC_RELS = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"><Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/styles" Target="styles.xml"/><Relationship Id="rId2" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes" Target="footnotes.xml"/><Relationship Id="rId3" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments" Target="comments.xml"/><Relationship Id="rId4" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/theme" Target="theme/theme1.xml"/><Relationship Id="rId8" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/footer" Target="footer1.xml"/><Relationship Id="rId9" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink" Target="https://example.com/" TargetMode="External"/></Relationships>'''

with zipfile.ZipFile(f"{out_dir}/pathological.docx", "w", zipfile.ZIP_DEFLATED) as z:
    z.writestr("[Content_Types].xml", CT)
    z.writestr("_rels/.rels", RELS)
    z.writestr("word/document.xml", DOC)
    z.writestr("word/_rels/document.xml.rels", DOC_RELS)
    z.writestr("word/styles.xml", STYLES)
    z.writestr("word/footnotes.xml", FOOTNOTES)
    z.writestr("word/comments.xml", COMMENTS)
    z.writestr("word/footer1.xml", FOOTER)
    z.writestr("word/theme/theme1.xml", THEME)
    z.writestr("word/media/blob.bin", b"\x00\x01\x02binary")
print("wrote pathological.docx")
