#!/usr/bin/env python3
"""Generate the round-trip corpus (SPEC §10.2): ~60 .docx files spanning
python-docx / Word / Google Docs / LibreOffice-style output plus deliberately
pathological cases.

    python3 tools/make_corpus.py corpus      # needs python-docx for the first set

Every file must pass crates/wp-docx/tests/roundtrip.rs. Files are named by
origin: gen-* (python-docx), word-*, gdocs-*, lo-*, path-* (hand-built XML).
"""
import struct
import sys
import zipfile
import zlib

out_dir = sys.argv[1] if len(sys.argv) > 1 else "corpus"
written = []


def save(name, z_entries, stored=(), order=None):
    """Write a .docx from {name: bytes|str}. `stored` entries are uncompressed."""
    names = order or list(z_entries)
    with zipfile.ZipFile(f"{out_dir}/{name}", "w") as z:
        for n in names:
            data = z_entries[n]
            if isinstance(data, str):
                data = data.encode("utf-8")
            z.writestr(n, data, zipfile.ZIP_STORED if n in stored else zipfile.ZIP_DEFLATED)
    written.append(name)


def png(w=8, h=8, rgb=(200, 30, 30)):
    raw = b"".join(b"\x00" + bytes(rgb) * w for _ in range(h))
    def chunk(t, d):
        return struct.pack(">I", len(d)) + t + d + struct.pack(">I", zlib.crc32(t + d) & 0xFFFFFFFF)
    return b"\x89PNG\r\n\x1a\n" + chunk(b"IHDR", struct.pack(">IIBBBBB", w, h, 8, 2, 0, 0, 0)) + chunk(b"IDAT", zlib.compress(raw)) + chunk(b"IEND", b"")


LOREM = ("Lorem ipsum dolor sit amet, consectetur adipiscing elit, sed do eiusmod tempor "
         "incididunt ut labore et dolore magna aliqua. Ut enim ad minim veniam, quis nostrud "
         "exercitation ullamco laboris nisi ut aliquip ex ea commodo consequat. ")

# ---------------------------------------------------------------------------
# python-docx set
# ---------------------------------------------------------------------------
try:
    import docx
    from docx.enum.section import WD_ORIENT, WD_SECTION
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER
    from docx.shared import Inches, Pt, RGBColor
    HAVE_DOCX = True
except ImportError:
    HAVE_DOCX = False
    print("python-docx not available; skipping gen-* fixtures")

if HAVE_DOCX:
    import io

    def gen(name):
        def deco(f):
            d = docx.Document()
            f(d)
            d.save(f"{out_dir}/{name}")
            written.append(name)
            return f
        return deco

    @gen("gen-report.docx")
    def _(d):
        d.add_heading("Quarterly Report", 0)
        d.add_heading("Summary", 1)
        p = d.add_paragraph("The quarterly results ")
        p.add_run("exceeded").bold = True
        p.add_run(" projections in every region except ")
        r = p.add_run("EMEA")
        r.italic = True
        r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
        p.add_run(", where currency effects reduced reported growth by roughly four points.")
        d.add_heading("Regional detail", 2)
        p = d.add_paragraph()
        r = p.add_run("North America grew 14% year over year. ")
        r.font.size = Pt(14)
        r = p.add_run("Underlined")
        r.underline = True
        p.add_run(" and ")
        r = p.add_run("struck")
        r.font.strike = True
        p.add_run(" and x")
        r = p.add_run("2")
        r.font.superscript = True
        p.add_run(".\tTabbed text here.")
        d.add_paragraph("First bullet", style="List Bullet")
        d.add_paragraph("Second bullet", style="List Bullet")
        d.add_paragraph("Numbered one", style="List Number")
        d.add_paragraph("Numbered two", style="List Number")
        p = d.add_paragraph("Centered paragraph with a hanging indent.")
        p.paragraph_format.alignment = 1
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.keep_with_next = True
        t = d.add_table(rows=2, cols=3)
        t.style = "Table Grid"
        for i in range(2):
            for j in range(3):
                t.cell(i, j).text = f"r{i}c{j}"
        p = d.add_paragraph("Text after the table, before a page break.")
        p.add_run().add_break(WD_BREAK.PAGE)
        d.add_paragraph("Appendix A — Methodology")
        for i in range(40):
            d.add_paragraph(LOREM * 2)

    @gen("gen-empty.docx")
    def _(d):
        pass

    @gen("gen-one-line.docx")
    def _(d):
        d.add_paragraph("Just one line.")

    @gen("gen-headings.docx")
    def _(d):
        for lvl in range(0, 10):
            d.add_heading(f"Heading level {lvl}", lvl)
            d.add_paragraph(LOREM)

    @gen("gen-char-formatting.docx")
    def _(d):
        p = d.add_paragraph()
        for attr in ["bold", "italic", "underline", "strike", "double_strike", "superscript", "subscript",
                     "small_caps", "all_caps", "shadow", "outline", "emboss", "imprint", "hidden", "highlight_color"]:
            r = p.add_run(attr + " ")
            if attr == "highlight_color":
                from docx.enum.text import WD_COLOR_INDEX
                r.font.highlight_color = WD_COLOR_INDEX.YELLOW
            else:
                setattr(r.font, attr, True)
        p = d.add_paragraph()
        for size in [6, 8, 9, 10, 10.5, 11, 12, 14, 18, 24, 36, 72]:
            r = p.add_run(f"{size}pt ")
            r.font.size = Pt(size)
        p = d.add_paragraph()
        for font in ["Calibri", "Cambria", "Arial", "Times New Roman", "Courier New", "Georgia", "Verdana", "Aptos", "Segoe UI", "Consolas"]:
            r = p.add_run(font + " ")
            r.font.name = font
        p = d.add_paragraph()
        for c in [(255, 0, 0), (0, 128, 0), (0, 0, 255), (128, 128, 128), (0, 0, 0)]:
            r = p.add_run("colour ")
            r.font.color.rgb = RGBColor(*c)
        r = d.add_paragraph().add_run("Nested: ")
        r.bold = True
        r.italic = True
        r.underline = True
        r.font.size = Pt(16)
        r.font.name = "Georgia"

    @gen("gen-para-formatting.docx")
    def _(d):
        for al in [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.JUSTIFY, WD_ALIGN_PARAGRAPH.DISTRIBUTE]:
            p = d.add_paragraph(LOREM)
            p.paragraph_format.alignment = al
        for li, ri, fi in [(0.5, 0, 0), (1, 1, 0.5), (0.5, 0, -0.5), (0, 0, 0.3), (2, 0.5, -1)]:
            p = d.add_paragraph(LOREM)
            p.paragraph_format.left_indent = Inches(li)
            p.paragraph_format.right_indent = Inches(ri)
            p.paragraph_format.first_line_indent = Inches(fi)
        for rule, val in [(WD_LINE_SPACING.SINGLE, None), (WD_LINE_SPACING.ONE_POINT_FIVE, None), (WD_LINE_SPACING.DOUBLE, None),
                          (WD_LINE_SPACING.EXACTLY, Pt(18)), (WD_LINE_SPACING.AT_LEAST, Pt(14)), (WD_LINE_SPACING.MULTIPLE, 1.15)]:
            p = d.add_paragraph(LOREM)
            if val is None:
                p.paragraph_format.line_spacing_rule = rule
            else:
                p.paragraph_format.line_spacing = val
        p = d.add_paragraph(LOREM)
        p.paragraph_format.space_before = Pt(24)
        p.paragraph_format.space_after = Pt(36)
        p.paragraph_format.keep_together = True
        p.paragraph_format.keep_with_next = True
        p.paragraph_format.widow_control = False
        p.paragraph_format.page_break_before = True
        p = d.add_paragraph("Tabs:\tleft\tcenter\tright\tdecimal 3.14")
        ts = p.paragraph_format.tab_stops
        ts.add_tab_stop(Inches(1))
        ts.add_tab_stop(Inches(2.5), WD_TAB_ALIGNMENT.CENTER)
        ts.add_tab_stop(Inches(4), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        ts.add_tab_stop(Inches(5.5), WD_TAB_ALIGNMENT.DECIMAL)

    @gen("gen-styles.docx")
    def _(d):
        for s in ["Normal", "No Spacing", "Title", "Subtitle", "Quote", "Intense Quote", "List Paragraph", "Caption",
                  "Body Text", "Body Text 2", "Body Text 3", "macro", "Heading 1", "Heading 2", "List", "List 2",
                  "List Bullet 2", "List Number 2", "List Continue", "TOC Heading"]:
            d.add_paragraph(f"Style: {s}. " + LOREM[:80], style=s)
        p = d.add_paragraph("Character styles: ")
        p.add_run("Strong", style="Strong")
        p.add_run(" and ")
        p.add_run("Emphasis", style="Emphasis")
        p.add_run(" and ")
        p.add_run("Intense Emphasis", style="Intense Emphasis")
        st = d.styles.add_style("My Custom Para", docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
        st.base_style = d.styles["Normal"]
        st.font.size = Pt(13)
        st.font.italic = True
        st.paragraph_format.left_indent = Inches(0.75)
        d.add_paragraph("A custom style based on Normal.", style="My Custom Para")

    @gen("gen-lists.docx")
    def _(d):
        for s in ["List Bullet", "List Bullet 2", "List Bullet 3", "List Number", "List Number 2", "List Number 3"]:
            for i in range(3):
                d.add_paragraph(f"{s} item {i + 1}", style=s)
        d.add_paragraph("Plain paragraph between lists.")
        for i in range(5):
            d.add_paragraph(f"Continued numbering {i + 1}", style="List Number")

    @gen("gen-tables.docx")
    def _(d):
        d.add_paragraph("Simple table:")
        t = d.add_table(rows=3, cols=3)
        t.style = "Table Grid"
        for i in range(3):
            for j in range(3):
                t.cell(i, j).text = f"cell {i},{j}"
        d.add_paragraph("Merged cells and a nested table:")
        t = d.add_table(rows=3, cols=4)
        t.style = "Light Grid Accent 1"
        a = t.cell(0, 0).merge(t.cell(0, 3))
        a.text = "Merged header across four columns"
        b = t.cell(1, 0).merge(t.cell(2, 0))
        b.text = "Merged down"
        inner = t.cell(1, 1).add_table(rows=2, cols=2)
        inner.cell(0, 0).text = "nested"
        p = t.cell(2, 3).paragraphs[0]
        r = p.add_run("bold in cell")
        r.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        d.add_paragraph("Two tables back to back:")
        d.add_table(rows=1, cols=1).cell(0, 0).text = "first"
        d.add_paragraph()
        d.add_table(rows=1, cols=1).cell(0, 0).text = "second"
        d.add_paragraph("Table at the very end.")
        t = d.add_table(rows=2, cols=2)
        t.cell(1, 1).text = "end"

    @gen("gen-sections.docx")
    def _(d):
        d.add_paragraph("Portrait Letter section." + LOREM)
        s = d.add_section(WD_SECTION.NEW_PAGE)
        s.orientation = WD_ORIENT.LANDSCAPE
        s.page_width, s.page_height = s.page_height, s.page_width
        d.add_paragraph("Landscape section." + LOREM)
        s = d.add_section(WD_SECTION.CONTINUOUS)
        s.left_margin = Inches(2)
        s.right_margin = Inches(2)
        d.add_paragraph("Continuous section with wide margins." + LOREM)
        s = d.add_section(WD_SECTION.ODD_PAGE)
        s.page_width = Inches(8.27)
        s.page_height = Inches(11.69)
        s.top_margin = Inches(0.5)
        d.add_paragraph("A4 odd-page section." + LOREM)

    @gen("gen-headers-footers.docx")
    def _(d):
        s = d.sections[0]
        s.header.paragraphs[0].text = "Header text"
        s.footer.paragraphs[0].text = "Footer text"
        s.different_first_page_header_footer = True
        s.first_page_header.paragraphs[0].text = "First page header"
        d.settings.odd_and_even_pages_header_footer = True
        s.even_page_header.paragraphs[0].text = "Even page header"
        for i in range(60):
            d.add_paragraph(LOREM)

    @gen("gen-images.docx")
    def _(d):
        d.add_paragraph("An inline picture:")
        d.add_picture(io.BytesIO(png()), width=Inches(1.5))
        p = d.add_paragraph("Picture inside a paragraph ")
        p.add_run().add_picture(io.BytesIO(png(4, 4, (0, 0, 200))), width=Inches(0.5))
        p.add_run(" with text after.")
        t = d.add_table(rows=1, cols=1)
        t.cell(0, 0).paragraphs[0].add_run().add_picture(io.BytesIO(png(6, 6, (0, 150, 0))), width=Inches(1))

    @gen("gen-breaks.docx")
    def _(d):
        p = d.add_paragraph("Line one")
        p.add_run().add_break(WD_BREAK.LINE)
        p.add_run("line two after a soft break")
        p.add_run().add_break(WD_BREAK.LINE_CLEAR_ALL)
        p.add_run("line three")
        p = d.add_paragraph("Before page break")
        p.add_run().add_break(WD_BREAK.PAGE)
        p.add_run("after page break in the same paragraph")
        d.add_page_break()
        d.add_paragraph("Paragraph after a paragraph-level page break")
        p = d.add_paragraph("Column break: ")
        p.add_run().add_break(WD_BREAK.COLUMN)
        p.add_run("after column break")

    @gen("gen-unicode.docx")
    def _(d):
        d.add_paragraph("Ünïcödé — “smart quotes” … ‘single’ – en dash — em dash • bullet © ® ™ € £ ¥")
        d.add_paragraph("日本語のテキスト。中文文本。한국어 텍스트。")
        d.add_paragraph("العربية والعبرية עברית (right-to-left scripts)")
        d.add_paragraph("Emoji: 😀 🎉 👍🏽 and combining: é (e + ́) ñ")
        d.add_paragraph("Math: ∑ ∫ √ ∞ ≠ ≤ ≥ α β γ δ ε")
        d.add_paragraph("Whitespace:  double  spaces nbsp thin​zwsp end")
        d.add_paragraph("Control-ish: ­ soft hyphen, ‑ no-break hyphen, tab\there")
        d.add_paragraph("<xml> & entities \" ' ]]> </w:t>")

    @gen("gen-long.docx")
    def _(d):
        d.add_heading("A long document", 0)
        for c in range(12):
            d.add_heading(f"Chapter {c + 1}", 1)
            for s in range(4):
                d.add_heading(f"Section {c + 1}.{s + 1}", 2)
                for _ in range(6):
                    p = d.add_paragraph(LOREM * 3)
                    p.add_run("Emphasised sentence at the end.").italic = True

    @gen("gen-core-props.docx")
    def _(d):
        cp = d.core_properties
        cp.author = "Ada Lovelace"
        cp.title = "Core properties test"
        cp.subject = "Metadata"
        cp.keywords = "wp, corpus, metadata"
        cp.comments = "Comments in docProps/core.xml"
        cp.category = "Test"
        cp.revision = 7
        d.add_paragraph("Document with core properties set.")

    @gen("gen-hyperlink-field.docx")
    def _(d):
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        p = d.add_paragraph("A hyperlink: ")
        part = d.part
        r_id = part.relate_to("https://example.org/page", docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
        h = OxmlElement("w:hyperlink")
        h.set(qn("r:id"), r_id)
        r = OxmlElement("w:r")
        rpr = OxmlElement("w:rPr")
        st = OxmlElement("w:rStyle")
        st.set(qn("w:val"), "Hyperlink")
        rpr.append(st)
        r.append(rpr)
        t = OxmlElement("w:t")
        t.text = "example.org"
        r.append(t)
        h.append(r)
        p._p.append(h)
        p = d.add_paragraph("Page ")
        for kind, text in [("begin", None), (None, " PAGE \\* MERGEFORMAT "), ("separate", None), (None, "1"), ("end", None)]:
            r = OxmlElement("w:r")
            if kind:
                fc = OxmlElement("w:fldChar")
                fc.set(qn("w:fldCharType"), kind)
                r.append(fc)
            elif text.strip().startswith("PAGE"):
                it = OxmlElement("w:instrText")
                it.set(qn("xml:space"), "preserve")
                it.text = text
                r.append(it)
            else:
                t = OxmlElement("w:t")
                t.text = text
                r.append(t)
            p._p.append(r)
        p.add_run(" of many.")

    @gen("gen-margins-paper.docx")
    def _(d):
        s = d.sections[0]
        s.page_width = Inches(5.5)
        s.page_height = Inches(8.5)
        s.left_margin = Inches(0.75)
        s.right_margin = Inches(0.75)
        s.top_margin = Inches(1.25)
        s.bottom_margin = Inches(0.6)
        s.gutter = Inches(0.25)
        s.header_distance = Inches(0.3)
        s.footer_distance = Inches(0.3)
        for _ in range(10):
            d.add_paragraph(LOREM)

    @gen("gen-empty-paragraphs.docx")
    def _(d):
        d.add_paragraph()
        d.add_paragraph()
        d.add_paragraph("Text between blank paragraphs")
        d.add_paragraph()
        p = d.add_paragraph()
        p.add_run("")
        p.add_run("").bold = True
        d.add_paragraph(" ")
        d.add_paragraph("\t")
        d.add_paragraph()

    @gen("gen-mixed-runs.docx")
    def _(d):
        p = d.add_paragraph()
        for i, word in enumerate((LOREM * 2).split()):
            r = p.add_run(word + " ")
            if i % 3 == 0:
                r.bold = True
            if i % 5 == 0:
                r.italic = True
            if i % 7 == 0:
                r.font.size = Pt(9)
            if i % 11 == 0:
                r.font.color.rgb = RGBColor(0, 0, 200)
        p = d.add_paragraph()
        for i in range(30):
            p.add_run("same ")  # many adjacent runs with identical formatting

# ---------------------------------------------------------------------------
# Hand-built XML set. Templates mimic real producers.
# ---------------------------------------------------------------------------
NS_WORD = ('xmlns:wpc="http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas" '
           'xmlns:cx="http://schemas.microsoft.com/office/drawing/2014/chartex" '
           'xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006" '
           'xmlns:o="urn:schemas-microsoft-com:office:office" '
           'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" '
           'xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math" '
           'xmlns:v="urn:schemas-microsoft-com:vml" '
           'xmlns:wp14="http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing" '
           'xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" '
           'xmlns:w10="urn:schemas-microsoft-com:office:word" '
           'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
           'xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml" '
           'xmlns:w15="http://schemas.microsoft.com/office/word/2012/wordml" '
           'xmlns:w16se="http://schemas.microsoft.com/office/word/2015/wordml/symex" '
           'xmlns:wpg="http://schemas.microsoft.com/office/word/2010/wordprocessingGroup" '
           'xmlns:wpi="http://schemas.microsoft.com/office/word/2010/wordprocessingInk" '
           'xmlns:wne="http://schemas.microsoft.com/office/word/2006/wordml" '
           'xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape" '
           'mc:Ignorable="w14 w15 w16se wp14"')
NS_MIN = ('xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
          'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"')
DECL = '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\r\n'
CT_MAIN = "application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"
CT = {
    "styles": "application/vnd.openxmlformats-officedocument.wordprocessingml.styles+xml",
    "numbering": "application/vnd.openxmlformats-officedocument.wordprocessingml.numbering+xml",
    "settings": "application/vnd.openxmlformats-officedocument.wordprocessingml.settings+xml",
    "webSettings": "application/vnd.openxmlformats-officedocument.wordprocessingml.webSettings+xml",
    "fontTable": "application/vnd.openxmlformats-officedocument.wordprocessingml.fontTable+xml",
    "footnotes": "application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml",
    "endnotes": "application/vnd.openxmlformats-officedocument.wordprocessingml.endnotes+xml",
    "comments": "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml",
    "header": "application/vnd.openxmlformats-officedocument.wordprocessingml.header+xml",
    "footer": "application/vnd.openxmlformats-officedocument.wordprocessingml.footer+xml",
    "theme": "application/vnd.openxmlformats-officedocument.theme+xml",
    "core": "application/vnd.openxmlformats-package.core-properties+xml",
    "app": "application/vnd.openxmlformats-officedocument.extended-properties+xml",
}
REL = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/"


def package(name, body, styles=None, extra=None, rels=None, ns=NS_WORD, decl=DECL, sect=None, stored=(), main="word/document.xml", pretty=False, order=None):
    """Build a package. `extra` = {part_path: (content_type_key, xml)}; `rels` = [(id, type_key, target, external)]."""
    extra = extra or {}
    rels = rels or []
    sect = sect if sect is not None else '<w:sectPr><w:pgSz w:w="12240" w:h="15840"/><w:pgMar w:top="1440" w:right="1440" w:bottom="1440" w:left="1440" w:header="720" w:footer="720" w:gutter="0"/><w:cols w:space="720"/><w:docGrid w:linePitch="360"/></w:sectPr>'
    doc = f'{decl}<w:document {ns}><w:body>{body}{sect}</w:body></w:document>'
    if pretty:
        doc = doc.replace("><", ">\n<")
    styles = styles if styles is not None else STYLES_WORD
    parts = {}
    overrides = [f'<Override PartName="/{main}" ContentType="{CT_MAIN}"/>']
    doc_rels = []
    mdir = main.rsplit("/", 1)[0]
    if styles:
        parts[f"{mdir}/styles.xml"] = styles
        overrides.append(f'<Override PartName="/{mdir}/styles.xml" ContentType="{CT["styles"]}"/>')
        doc_rels.append(f'<Relationship Id="rId1" Type="{REL}styles" Target="styles.xml"/>')
    for path, (ct, xml) in extra.items():
        parts[path] = xml
        if ct in CT:
            overrides.append(f'<Override PartName="/{path}" ContentType="{CT[ct]}"/>')
    for rid, ty, target, ext in rels:
        mode = ' TargetMode="External"' if ext else ""
        doc_rels.append(f'<Relationship Id="{rid}" Type="{REL}{ty}" Target="{target}"{mode}/>')
    parts["[Content_Types].xml"] = (decl + '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
                                    '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
                                    '<Default Extension="xml" ContentType="application/xml"/>'
                                    '<Default Extension="png" ContentType="image/png"/>' + "".join(overrides) + '</Types>')
    parts["_rels/.rels"] = (decl + '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
                            f'<Relationship Id="rId1" Type="{REL}officeDocument" Target="{main}"/></Relationships>')
    parts[f"{mdir}/_rels/{main.rsplit('/', 1)[1]}.rels"] = (decl + '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">' + "".join(doc_rels) + '</Relationships>')
    parts[main] = doc
    save(name, parts, stored=stored, order=order)


def p(text, ppr="", rpr="", attrs=' w:rsidR="00C62A19" w:rsidRDefault="00C62A19" w14:paraId="1A2B3C4D" w14:textId="77777777"', run_attrs=' w:rsidRPr="00C62A19"'):
    pp = f"<w:pPr>{ppr}</w:pPr>" if ppr else ""
    rp = f"<w:rPr>{rpr}</w:rPr>" if rpr else ""
    if text == "":
        return f"<w:p{attrs}>{pp}</w:p>"
    return f'<w:p{attrs}>{pp}<w:r{run_attrs}>{rp}<w:t xml:space="preserve">{text}</w:t></w:r></w:p>'


def pw(*runs, ppr="", attrs=' w:rsidR="00C62A19" w:rsidRDefault="00C62A19"'):
    """Paragraph from (text, rpr) runs, Word-style with proofErr and rsids."""
    pp = f"<w:pPr>{ppr}</w:pPr>" if ppr else ""
    out = [f"<w:p{attrs}>{pp}"]
    for i, (text, rpr) in enumerate(runs):
        rp = f"<w:rPr>{rpr}</w:rPr>" if rpr else ""
        if i == 1:
            out.append('<w:proofErr w:type="spellStart"/>')
        out.append(f'<w:r w:rsidR="00{i:02d}AB12">{rp}<w:t xml:space="preserve">{text}</w:t></w:r>')
        if i == 1:
            out.append('<w:proofErr w:type="spellEnd"/>')
    out.append("</w:p>")
    return "".join(out)


STYLES_WORD = f'''{DECL}<w:styles {NS_MIN} xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006" xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml" mc:Ignorable="w14"><w:docDefaults><w:rPrDefault><w:rPr><w:rFonts w:asciiTheme="minorHAnsi" w:eastAsiaTheme="minorHAnsi" w:hAnsiTheme="minorHAnsi" w:cstheme="minorBidi"/><w:sz w:val="22"/><w:szCs w:val="22"/><w:lang w:val="en-US" w:eastAsia="en-US" w:bidi="ar-SA"/></w:rPr></w:rPrDefault><w:pPrDefault><w:pPr><w:spacing w:after="160" w:line="259" w:lineRule="auto"/></w:pPr></w:pPrDefault></w:docDefaults><w:latentStyles w:defLockedState="0" w:defUIPriority="99" w:defSemiHidden="0" w:defUnhideWhenUsed="0" w:defQFormat="0" w:count="376"><w:lsdException w:name="Normal" w:uiPriority="0" w:qFormat="1"/><w:lsdException w:name="heading 1" w:uiPriority="9" w:qFormat="1"/></w:latentStyles><w:style w:type="paragraph" w:default="1" w:styleId="Normal"><w:name w:val="Normal"/><w:qFormat/></w:style><w:style w:type="paragraph" w:styleId="Heading1"><w:name w:val="heading 1"/><w:basedOn w:val="Normal"/><w:next w:val="Normal"/><w:link w:val="Heading1Char"/><w:uiPriority w:val="9"/><w:qFormat/><w:rsid w:val="00C62A19"/><w:pPr><w:keepNext/><w:keepLines/><w:spacing w:before="240" w:after="0"/><w:outlineLvl w:val="0"/></w:pPr><w:rPr><w:rFonts w:asciiTheme="majorHAnsi" w:eastAsiaTheme="majorHAnsi" w:hAnsiTheme="majorHAnsi" w:cstheme="majorBidi"/><w:color w:val="2F5496" w:themeColor="accent1" w:themeShade="BF"/><w:sz w:val="32"/><w:szCs w:val="32"/></w:rPr></w:style><w:style w:type="paragraph" w:styleId="Heading2"><w:name w:val="heading 2"/><w:basedOn w:val="Normal"/><w:next w:val="Normal"/><w:uiPriority w:val="9"/><w:unhideWhenUsed/><w:qFormat/><w:pPr><w:keepNext/><w:keepLines/><w:spacing w:before="40" w:after="0"/><w:outlineLvl w:val="1"/></w:pPr><w:rPr><w:rFonts w:asciiTheme="majorHAnsi" w:hAnsiTheme="majorHAnsi"/><w:color w:val="2F5496"/><w:sz w:val="26"/></w:rPr></w:style><w:style w:type="character" w:default="1" w:styleId="DefaultParagraphFont"><w:name w:val="Default Paragraph Font"/><w:uiPriority w:val="1"/><w:semiHidden/><w:unhideWhenUsed/></w:style><w:style w:type="table" w:default="1" w:styleId="TableNormal"><w:name w:val="Normal Table"/><w:uiPriority w:val="99"/><w:semiHidden/><w:unhideWhenUsed/><w:tblPr><w:tblInd w:w="0" w:type="dxa"/><w:tblCellMar><w:top w:w="0" w:type="dxa"/><w:left w:w="108" w:type="dxa"/><w:bottom w:w="0" w:type="dxa"/><w:right w:w="108" w:type="dxa"/></w:tblCellMar></w:tblPr></w:style><w:style w:type="numbering" w:default="1" w:styleId="NoList"><w:name w:val="No List"/><w:uiPriority w:val="99"/><w:semiHidden/><w:unhideWhenUsed/></w:style><w:style w:type="character" w:customStyle="1" w:styleId="Heading1Char"><w:name w:val="Heading 1 Char"/><w:basedOn w:val="DefaultParagraphFont"/><w:link w:val="Heading1"/><w:uiPriority w:val="9"/><w:rPr><w:rFonts w:asciiTheme="majorHAnsi" w:hAnsiTheme="majorHAnsi"/><w:color w:val="2F5496"/><w:sz w:val="32"/></w:rPr></w:style><w:style w:type="paragraph" w:styleId="ListParagraph"><w:name w:val="List Paragraph"/><w:basedOn w:val="Normal"/><w:uiPriority w:val="34"/><w:qFormat/><w:pPr><w:ind w:left="720"/><w:contextualSpacing/></w:pPr></w:style><w:style w:type="character" w:styleId="Hyperlink"><w:name w:val="Hyperlink"/><w:basedOn w:val="DefaultParagraphFont"/><w:uiPriority w:val="99"/><w:unhideWhenUsed/><w:rPr><w:color w:val="0563C1" w:themeColor="hyperlink"/><w:u w:val="single"/></w:rPr></w:style><w:style w:type="table" w:styleId="TableGrid"><w:name w:val="Table Grid"/><w:basedOn w:val="TableNormal"/><w:uiPriority w:val="39"/><w:pPr><w:spacing w:after="0" w:line="240" w:lineRule="auto"/></w:pPr><w:tblPr><w:tblBorders><w:top w:val="single" w:sz="4" w:space="0" w:color="auto"/><w:left w:val="single" w:sz="4" w:space="0" w:color="auto"/><w:bottom w:val="single" w:sz="4" w:space="0" w:color="auto"/><w:right w:val="single" w:sz="4" w:space="0" w:color="auto"/><w:insideH w:val="single" w:sz="4" w:space="0" w:color="auto"/><w:insideV w:val="single" w:sz="4" w:space="0" w:color="auto"/></w:tblBorders></w:tblPr></w:style><w:style w:type="paragraph" w:styleId="FootnoteText"><w:name w:val="footnote text"/><w:basedOn w:val="Normal"/><w:pPr><w:spacing w:after="0" w:line="240" w:lineRule="auto"/></w:pPr><w:rPr><w:sz w:val="20"/></w:rPr></w:style><w:style w:type="character" w:styleId="FootnoteReference"><w:name w:val="footnote reference"/><w:rPr><w:vertAlign w:val="superscript"/></w:rPr></w:style></w:styles>'''

NUMBERING_WORD = f'''{DECL}<w:numbering {NS_MIN} xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml"><w:abstractNum w:abstractNumId="0" w15:restartNumberingAfterBreak="0" xmlns:w15="http://schemas.microsoft.com/office/word/2012/wordml"><w:nsid w:val="0B8A1E3C"/><w:multiLevelType w:val="hybridMultilevel"/><w:tmpl w:val="8F3A6C54"/><w:lvl w:ilvl="0" w:tplc="04090001"><w:start w:val="1"/><w:numFmt w:val="bullet"/><w:lvlText w:val="&#61623;"/><w:lvlJc w:val="left"/><w:pPr><w:ind w:left="720" w:hanging="360"/></w:pPr><w:rPr><w:rFonts w:ascii="Symbol" w:hAnsi="Symbol" w:hint="default"/></w:rPr></w:lvl><w:lvl w:ilvl="1" w:tplc="04090003"><w:start w:val="1"/><w:numFmt w:val="bullet"/><w:lvlText w:val="o"/><w:lvlJc w:val="left"/><w:pPr><w:ind w:left="1440" w:hanging="360"/></w:pPr><w:rPr><w:rFonts w:ascii="Courier New" w:hAnsi="Courier New" w:cs="Courier New" w:hint="default"/></w:rPr></w:lvl><w:lvl w:ilvl="2" w:tplc="04090005"><w:start w:val="1"/><w:numFmt w:val="bullet"/><w:lvlText w:val="&#61607;"/><w:lvlJc w:val="left"/><w:pPr><w:ind w:left="2160" w:hanging="360"/></w:pPr><w:rPr><w:rFonts w:ascii="Wingdings" w:hAnsi="Wingdings" w:hint="default"/></w:rPr></w:lvl></w:abstractNum><w:abstractNum w:abstractNumId="1"><w:nsid w:val="2C3D4E5F"/><w:multiLevelType w:val="multilevel"/><w:lvl w:ilvl="0"><w:start w:val="1"/><w:numFmt w:val="decimal"/><w:lvlText w:val="%1."/><w:lvlJc w:val="left"/><w:pPr><w:ind w:left="720" w:hanging="360"/></w:pPr></w:lvl><w:lvl w:ilvl="1"><w:start w:val="1"/><w:numFmt w:val="lowerLetter"/><w:lvlText w:val="%2)"/><w:lvlJc w:val="left"/><w:pPr><w:ind w:left="1440" w:hanging="360"/></w:pPr></w:lvl><w:lvl w:ilvl="2"><w:start w:val="1"/><w:numFmt w:val="lowerRoman"/><w:lvlText w:val="%3."/><w:lvlJc w:val="right"/><w:pPr><w:ind w:left="2160" w:hanging="180"/></w:pPr></w:lvl><w:lvl w:ilvl="3"><w:start w:val="1"/><w:numFmt w:val="decimal"/><w:lvlText w:val="%1.%2.%3.%4"/><w:lvlJc w:val="left"/><w:pPr><w:ind w:left="2880" w:hanging="360"/></w:pPr></w:lvl></w:abstractNum><w:abstractNum w:abstractNumId="2"><w:multiLevelType w:val="multilevel"/><w:lvl w:ilvl="0"><w:start w:val="1"/><w:numFmt w:val="upperRoman"/><w:lvlText w:val="%1."/><w:lvlJc w:val="left"/><w:pPr><w:ind w:left="360" w:hanging="360"/></w:pPr></w:lvl><w:lvl w:ilvl="1"><w:start w:val="1"/><w:numFmt w:val="upperLetter"/><w:lvlText w:val="%2."/><w:lvlJc w:val="left"/><w:pPr><w:ind w:left="720" w:hanging="360"/></w:pPr></w:lvl><w:lvl w:ilvl="2"><w:start w:val="1"/><w:numFmt w:val="decimalZero"/><w:lvlText w:val="%3"/><w:lvlJc w:val="left"/><w:pPr><w:ind w:left="1080" w:hanging="360"/></w:pPr></w:lvl><w:lvl w:ilvl="3"><w:start w:val="5"/><w:numFmt w:val="ordinal"/><w:lvlText w:val="%4"/><w:lvlJc w:val="left"/><w:pPr><w:ind w:left="1440" w:hanging="360"/></w:pPr></w:lvl><w:lvl w:ilvl="4"><w:start w:val="1"/><w:numFmt w:val="cardinalText"/><w:lvlText w:val="%5"/><w:lvlJc w:val="left"/><w:pPr><w:ind w:left="1800" w:hanging="360"/></w:pPr></w:lvl><w:lvl w:ilvl="5"><w:start w:val="1"/><w:numFmt w:val="none"/><w:lvlText w:val="-"/><w:lvlJc w:val="left"/><w:pPr><w:ind w:left="2160" w:hanging="360"/></w:pPr></w:lvl></w:abstractNum><w:num w:numId="1" w16cid:durableId="12345" xmlns:w16cid="http://schemas.microsoft.com/office/word/2016/wordml/cid"><w:abstractNumId w:val="0"/></w:num><w:num w:numId="2"><w:abstractNumId w:val="1"/></w:num><w:num w:numId="3"><w:abstractNumId w:val="1"/><w:lvlOverride w:ilvl="0"><w:startOverride w:val="1"/></w:lvlOverride></w:num><w:num w:numId="4"><w:abstractNumId w:val="2"/></w:num><w:num w:numId="5"><w:abstractNumId w:val="1"/><w:lvlOverride w:ilvl="0"><w:startOverride w:val="7"/></w:lvlOverride><w:lvlOverride w:ilvl="1"><w:lvl w:ilvl="1"><w:start w:val="1"/><w:numFmt w:val="upperLetter"/><w:lvlText w:val="(%2)"/><w:lvlJc w:val="left"/><w:pPr><w:ind w:left="1440" w:hanging="360"/></w:pPr></w:lvl></w:lvlOverride></w:num></w:numbering>'''

SETTINGS_WORD = f'{DECL}<w:settings {NS_MIN} xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml"><w:zoom w:percent="100"/><w:proofState w:spelling="clean" w:grammar="clean"/><w:defaultTabStop w:val="720"/><w:characterSpacingControl w:val="doNotCompress"/><w:footnotePr><w:footnote w:id="-1"/><w:footnote w:id="0"/></w:footnotePr><w:compat><w:compatSetting w:name="compatibilityMode" w:uri="http://schemas.microsoft.com/office/word" w:val="15"/></w:compat><w:rsids><w:rsidRoot w:val="00C62A19"/><w:rsid w:val="00C62A19"/><w:rsid w:val="0001AB12"/></w:rsids><m:mathPr xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"><m:mathFont m:val="Cambria Math"/></m:mathPr><w:themeFontLang w:val="en-US"/><w:decimalSymbol w:val="."/><w:listSeparator w:val=","/><w14:docId w14:val="4D4B1F3E"/></w:settings>'
THEME = f'{DECL}<a:theme xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" name="Office Theme"><a:themeElements><a:clrScheme name="Office"><a:dk1><a:sysClr val="windowText" lastClr="000000"/></a:dk1><a:lt1><a:sysClr val="window" lastClr="FFFFFF"/></a:lt1><a:accent1><a:srgbClr val="4472C4"/></a:accent1><a:hlink><a:srgbClr val="0563C1"/></a:hlink></a:clrScheme><a:fontScheme name="Office"><a:majorFont><a:latin typeface="Calibri Light" panose="020F0302020204030204"/><a:ea typeface=""/><a:cs typeface=""/></a:majorFont><a:minorFont><a:latin typeface="Calibri" panose="020F0502020204030204"/><a:ea typeface=""/><a:cs typeface=""/></a:minorFont></a:fontScheme></a:themeElements></a:theme>'
FOOTNOTES = f'{DECL}<w:footnotes {NS_MIN}><w:footnote w:type="separator" w:id="-1"><w:p><w:pPr><w:spacing w:after="0" w:line="240" w:lineRule="auto"/></w:pPr><w:r><w:separator/></w:r></w:p></w:footnote><w:footnote w:type="continuationSeparator" w:id="0"><w:p><w:r><w:continuationSeparator/></w:r></w:p></w:footnote><w:footnote w:id="1"><w:p><w:pPr><w:pStyle w:val="FootnoteText"/></w:pPr><w:r><w:rPr><w:rStyle w:val="FootnoteReference"/></w:rPr><w:footnoteRef/></w:r><w:r><w:t xml:space="preserve"> The first footnote, with </w:t></w:r><w:r><w:rPr><w:i/></w:rPr><w:t>italics</w:t></w:r><w:r><w:t>.</w:t></w:r></w:p></w:footnote><w:footnote w:id="2"><w:p><w:pPr><w:pStyle w:val="FootnoteText"/></w:pPr><w:r><w:rPr><w:rStyle w:val="FootnoteReference"/></w:rPr><w:footnoteRef/></w:r><w:r><w:t xml:space="preserve"> Second footnote.</w:t></w:r></w:p></w:footnote></w:footnotes>'
ENDNOTES = f'{DECL}<w:endnotes {NS_MIN}><w:endnote w:type="separator" w:id="-1"><w:p><w:r><w:separator/></w:r></w:p></w:endnote><w:endnote w:id="1"><w:p><w:r><w:endnoteRef/></w:r><w:r><w:t xml:space="preserve"> An endnote.</w:t></w:r></w:p></w:endnote></w:endnotes>'
COMMENTS = f'{DECL}<w:comments {NS_MIN}><w:comment w:id="0" w:author="Reviewer One" w:date="2025-03-01T10:00:00Z" w:initials="R1"><w:p><w:pPr><w:pStyle w:val="CommentText"/></w:pPr><w:r><w:rPr><w:rStyle w:val="CommentReference"/></w:rPr><w:annotationRef/></w:r><w:r><w:t>Please check this figure.</w:t></w:r></w:p></w:comment><w:comment w:id="1" w:author="Reviewer Two" w:date="2025-03-02T10:00:00Z" w:initials="R2"><w:p><w:r><w:annotationRef/></w:r><w:r><w:t>Agreed.</w:t></w:r></w:p></w:comment></w:comments>'
HEADER = f'{DECL}<w:hdr {NS_MIN}><w:p><w:pPr><w:pStyle w:val="Header"/><w:jc w:val="right"/></w:pPr><w:r><w:t>Confidential draft</w:t></w:r></w:p></w:hdr>'
FOOTER = f'{DECL}<w:ftr {NS_MIN}><w:p><w:pPr><w:pStyle w:val="Footer"/><w:jc w:val="center"/></w:pPr><w:r><w:fldChar w:fldCharType="begin"/></w:r><w:r><w:instrText xml:space="preserve"> PAGE </w:instrText></w:r><w:r><w:fldChar w:fldCharType="separate"/></w:r><w:r><w:t>1</w:t></w:r><w:r><w:fldChar w:fldCharType="end"/></w:r></w:p></w:ftr>'
CORE = f'{DECL}<cp:coreProperties xmlns:cp="http://schemas.openxmlformats.org/package/2006/metadata/core-properties" xmlns:dc="http://purl.org/dc/elements/1.1/" xmlns:dcterms="http://purl.org/dc/terms/" xmlns:dcmitype="http://purl.org/dc/dcmitype/" xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance"><dc:title>Corpus</dc:title><dc:creator>Author</dc:creator><cp:lastModifiedBy>Author</cp:lastModifiedBy><cp:revision>3</cp:revision><dcterms:created xsi:type="dcterms:W3CDTF">2025-01-01T00:00:00Z</dcterms:created><dcterms:modified xsi:type="dcterms:W3CDTF">2025-06-01T12:34:56Z</dcterms:modified></cp:coreProperties>'
APP = f'{DECL}<Properties xmlns="http://schemas.openxmlformats.org/officeDocument/2006/extended-properties" xmlns:vt="http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes"><Template>Normal.dotm</Template><TotalTime>12</TotalTime><Pages>3</Pages><Words>321</Words><Application>Microsoft Office Word</Application><AppVersion>16.0000</AppVersion></Properties>'

WORD_EXTRA = {
    "word/settings.xml": ("settings", SETTINGS_WORD),
    "word/theme/theme1.xml": ("theme", THEME),
    "docProps/core.xml": ("core", CORE),
    "docProps/app.xml": ("app", APP),
}
WORD_RELS = [("rId2", "settings", "settings.xml", False), ("rId3", "theme", "theme/theme1.xml", False)]

# --- Word-style ------------------------------------------------------------
B = "<w:b/><w:bCs/>"
I = "<w:i/><w:iCs/>"

package("word-basic.docx",
        p("Title of the document", '<w:pStyle w:val="Heading1"/>')
        + pw(("The quick brown ", ""), ("fox", B), (" jumps over the lazy dog. ", ""), ("Lazy", I), (".", ""))
        + p("")
        + p("A second paragraph with rsids and a paragraph id.", '<w:spacing w:after="0"/><w:jc w:val="both"/>'),
        extra=WORD_EXTRA, rels=WORD_RELS)

package("word-lists.docx",
        p("Bulleted:", "")
        + p("First bullet", '<w:pStyle w:val="ListParagraph"/><w:numPr><w:ilvl w:val="0"/><w:numId w:val="1"/></w:numPr>')
        + p("Nested bullet", '<w:pStyle w:val="ListParagraph"/><w:numPr><w:ilvl w:val="1"/><w:numId w:val="1"/></w:numPr>')
        + p("Deeper bullet", '<w:pStyle w:val="ListParagraph"/><w:numPr><w:ilvl w:val="2"/><w:numId w:val="1"/></w:numPr>')
        + p("Back to top level", '<w:pStyle w:val="ListParagraph"/><w:numPr><w:ilvl w:val="0"/><w:numId w:val="1"/></w:numPr>')
        + p("Numbered:", "")
        + p("One", '<w:pStyle w:val="ListParagraph"/><w:numPr><w:ilvl w:val="0"/><w:numId w:val="2"/></w:numPr>')
        + p("Two", '<w:pStyle w:val="ListParagraph"/><w:numPr><w:ilvl w:val="0"/><w:numId w:val="2"/></w:numPr>')
        + p("Two a", '<w:pStyle w:val="ListParagraph"/><w:numPr><w:ilvl w:val="1"/><w:numId w:val="2"/></w:numPr>')
        + p("Two b", '<w:pStyle w:val="ListParagraph"/><w:numPr><w:ilvl w:val="1"/><w:numId w:val="2"/></w:numPr>')
        + p("Two b i", '<w:pStyle w:val="ListParagraph"/><w:numPr><w:ilvl w:val="2"/><w:numId w:val="2"/></w:numPr>')
        + p("Level four 2.2.1.1", '<w:pStyle w:val="ListParagraph"/><w:numPr><w:ilvl w:val="3"/><w:numId w:val="2"/></w:numPr>')
        + p("Three", '<w:pStyle w:val="ListParagraph"/><w:numPr><w:ilvl w:val="0"/><w:numId w:val="2"/></w:numPr>')
        + p("Interrupting paragraph.", "")
        + p("Restarted at one", '<w:pStyle w:val="ListParagraph"/><w:numPr><w:ilvl w:val="0"/><w:numId w:val="3"/></w:numPr>')
        + p("Roman I", '<w:numPr><w:ilvl w:val="0"/><w:numId w:val="4"/></w:numPr>')
        + p("Roman I.A", '<w:numPr><w:ilvl w:val="1"/><w:numId w:val="4"/></w:numPr>')
        + p("Zero-padded 01", '<w:numPr><w:ilvl w:val="2"/><w:numId w:val="4"/></w:numPr>')
        + p("Ordinal 5th", '<w:numPr><w:ilvl w:val="3"/><w:numId w:val="4"/></w:numPr>')
        + p("Cardinal text", '<w:numPr><w:ilvl w:val="4"/><w:numId w:val="4"/></w:numPr>')
        + p("No number, dash", '<w:numPr><w:ilvl w:val="5"/><w:numId w:val="4"/></w:numPr>')
        + p("Override start at seven", '<w:numPr><w:ilvl w:val="0"/><w:numId w:val="5"/></w:numPr>')
        + p("Override level (A)", '<w:numPr><w:ilvl w:val="1"/><w:numId w:val="5"/></w:numPr>')
        + p("List item with direct indent", '<w:numPr><w:ilvl w:val="0"/><w:numId w:val="2"/></w:numPr><w:ind w:left="1800" w:hanging="720"/>')
        + p("numId 0 removes numbering", '<w:pStyle w:val="ListParagraph"/><w:numPr><w:ilvl w:val="0"/><w:numId w:val="0"/></w:numPr>'),
        extra={**WORD_EXTRA, "word/numbering.xml": ("numbering", NUMBERING_WORD)},
        rels=WORD_RELS + [("rId4", "numbering", "numbering.xml", False)])

package("word-tracked-changes.docx",
        p("Tracked changes", '<w:pStyle w:val="Heading1"/>')
        + '<w:p w:rsidR="00A1" w:rsidDel="00A2"><w:r><w:t xml:space="preserve">Kept text </w:t></w:r><w:ins w:id="1" w:author="Ann" w:date="2025-01-01T09:00:00Z"><w:r><w:t>inserted text</w:t></w:r></w:ins><w:r><w:t xml:space="preserve"> and </w:t></w:r><w:del w:id="2" w:author="Ann" w:date="2025-01-01T09:00:00Z"><w:r><w:delText xml:space="preserve">deleted text </w:delText></w:r></w:del><w:r><w:t>end.</w:t></w:r></w:p>'
        + '<w:p><w:pPr><w:rPr><w:ins w:id="3" w:author="Ann" w:date="2025-01-01T09:00:00Z"/></w:rPr></w:pPr><w:ins w:id="4" w:author="Ann" w:date="2025-01-01T09:00:00Z"><w:r><w:t>An entirely inserted paragraph.</w:t></w:r></w:ins></w:p>'
        + '<w:p><w:pPr><w:rPr><w:del w:id="5" w:author="Bob" w:date="2025-01-02T09:00:00Z"/></w:rPr></w:pPr><w:del w:id="6" w:author="Bob" w:date="2025-01-02T09:00:00Z"><w:r><w:delText>An entirely deleted paragraph.</w:delText></w:r></w:del></w:p>'
        + '<w:p><w:r><w:rPr><w:rPrChange w:id="7" w:author="Bob" w:date="2025-01-02T09:00:00Z"><w:rPr/></w:rPrChange><w:b/></w:rPr><w:t>Formatting change to bold</w:t></w:r></w:p>'
        + '<w:p><w:pPr><w:jc w:val="center"/><w:pPrChange w:id="8" w:author="Bob" w:date="2025-01-02T09:00:00Z"><w:pPr/></w:pPrChange></w:pPr><w:r><w:t>Paragraph property change</w:t></w:r></w:p>'
        + '<w:p><w:moveFromRangeStart w:id="9" w:author="Ann" w:date="2025-01-03T09:00:00Z" w:name="move1"/><w:moveFrom w:id="10" w:author="Ann" w:date="2025-01-03T09:00:00Z"><w:r><w:delText>Moved sentence.</w:delText></w:r></w:moveFrom><w:moveFromRangeEnd w:id="9"/></w:p>'
        + '<w:p><w:moveToRangeStart w:id="11" w:author="Ann" w:date="2025-01-03T09:00:00Z" w:name="move1"/><w:moveTo w:id="12" w:author="Ann" w:date="2025-01-03T09:00:00Z"><w:r><w:t>Moved sentence.</w:t></w:r></w:moveTo><w:moveToRangeEnd w:id="11"/></w:p>'
        + '<w:p><w:ins w:id="13" w:author="Ann" w:date="2025-01-01T09:00:00Z"><w:del w:id="14" w:author="Bob" w:date="2025-01-02T09:00:00Z"><w:r><w:delText>inserted then deleted</w:delText></w:r></w:del></w:ins></w:p>',
        extra=WORD_EXTRA, rels=WORD_RELS)

package("word-comments.docx",
        p("Comments", '<w:pStyle w:val="Heading1"/>')
        + '<w:p><w:r><w:t xml:space="preserve">Some </w:t></w:r><w:commentRangeStart w:id="0"/><w:r><w:t>commented text</w:t></w:r><w:commentRangeEnd w:id="0"/><w:r><w:rPr><w:rStyle w:val="CommentReference"/></w:rPr><w:commentReference w:id="0"/></w:r><w:r><w:t xml:space="preserve"> and </w:t></w:r><w:commentRangeStart w:id="1"/><w:r><w:t>more</w:t></w:r><w:commentRangeEnd w:id="1"/><w:r><w:commentReference w:id="1"/></w:r><w:r><w:t>.</w:t></w:r></w:p>'
        + '<w:commentRangeStart w:id="2"/><w:p><w:r><w:t>Whole paragraph commented</w:t></w:r></w:p><w:commentRangeEnd w:id="2"/><w:p><w:r><w:commentReference w:id="2"/></w:r></w:p>',
        extra={**WORD_EXTRA, "word/comments.xml": ("comments", COMMENTS.replace("</w:comments>", '<w:comment w:id="2" w:author="X" w:date="2025-03-03T00:00:00Z"><w:p><w:r><w:t>Third.</w:t></w:r></w:p></w:comment></w:comments>'))},
        rels=WORD_RELS + [("rId5", "comments", "comments.xml", False)])

package("word-footnotes-endnotes.docx",
        p("Notes", '<w:pStyle w:val="Heading1"/>')
        + '<w:p><w:r><w:t>Text with a footnote</w:t></w:r><w:r><w:rPr><w:rStyle w:val="FootnoteReference"/></w:rPr><w:footnoteReference w:id="1"/></w:r><w:r><w:t xml:space="preserve"> and another</w:t></w:r><w:r><w:rPr><w:vertAlign w:val="superscript"/></w:rPr><w:footnoteReference w:id="2"/></w:r><w:r><w:t xml:space="preserve"> and an endnote</w:t></w:r><w:r><w:rPr><w:vertAlign w:val="superscript"/></w:rPr><w:endnoteReference w:id="1"/></w:r><w:r><w:t>.</w:t></w:r></w:p>'
        + '<w:p><w:r><w:t>Custom mark</w:t></w:r><w:r><w:footnoteReference w:customMarkFollows="1" w:id="2"/><w:t>*</w:t></w:r></w:p>',
        extra={**WORD_EXTRA, "word/footnotes.xml": ("footnotes", FOOTNOTES), "word/endnotes.xml": ("endnotes", ENDNOTES)},
        rels=WORD_RELS + [("rId6", "footnotes", "footnotes.xml", False), ("rId7", "endnotes", "endnotes.xml", False)])

package("word-headers-footers.docx",
        "".join(p(f"Paragraph {i}. " + LOREM) for i in range(30)),
        sect='<w:sectPr w:rsidR="00C62A19"><w:headerReference w:type="default" r:id="rId8"/><w:footerReference w:type="default" r:id="rId9"/><w:headerReference w:type="first" r:id="rId8"/><w:type w:val="nextPage"/><w:pgSz w:w="12240" w:h="15840"/><w:pgMar w:top="1440" w:right="1440" w:bottom="1440" w:left="1440" w:header="720" w:footer="720" w:gutter="0"/><w:pgNumType w:start="3"/><w:cols w:space="720"/><w:titlePg/><w:docGrid w:linePitch="360"/></w:sectPr>',
        extra={**WORD_EXTRA, "word/header1.xml": ("header", HEADER), "word/footer1.xml": ("footer", FOOTER)},
        rels=WORD_RELS + [("rId8", "header", "header1.xml", False), ("rId9", "footer", "footer1.xml", False)])

package("word-sections.docx",
        p("Section one, portrait.")
        + '<w:p><w:pPr><w:sectPr w:rsidR="00C62A19"><w:pgSz w:w="12240" w:h="15840"/><w:pgMar w:top="1440" w:right="1440" w:bottom="1440" w:left="1440" w:header="720" w:footer="720" w:gutter="0"/><w:cols w:space="720"/><w:docGrid w:linePitch="360"/></w:sectPr></w:pPr></w:p>'
        + p("Section two, landscape, two columns. " + LOREM * 3)
        + '<w:p><w:r><w:br w:type="column"/></w:r><w:r><w:t>Second column</w:t></w:r></w:p>'
        + '<w:p><w:pPr><w:sectPr><w:type w:val="continuous"/><w:pgSz w:w="15840" w:h="12240" w:orient="landscape"/><w:pgMar w:top="1440" w:right="1440" w:bottom="1440" w:left="1440" w:header="720" w:footer="720" w:gutter="0"/><w:cols w:num="2" w:space="720" w:equalWidth="0"><w:col w:w="4000" w:space="720"/><w:col w:w="8240"/></w:cols><w:docGrid w:linePitch="360"/></w:sectPr></w:pPr></w:p>'
        + p("Section three, back to portrait, A4."),
        sect='<w:sectPr><w:pgSz w:w="11906" w:h="16838" w:code="9"/><w:pgMar w:top="1134" w:right="1134" w:bottom="1134" w:left="1134" w:header="709" w:footer="709" w:gutter="0"/><w:lnNumType w:countBy="1" w:restart="continuous"/><w:cols w:space="708"/><w:docGrid w:linePitch="360"/></w:sectPr>',
        extra=WORD_EXTRA, rels=WORD_RELS)

DRAWING = ('<w:drawing><wp:inline distT="0" distB="0" distL="0" distR="0" wp14:anchorId="1F2E3D4C" wp14:editId="5A6B7C8D"><wp:extent cx="914400" cy="914400"/><wp:effectExtent l="0" t="0" r="0" b="0"/><wp:docPr id="1" name="Picture 1" descr="A red square"/><wp:cNvGraphicFramePr><a:graphicFrameLocks xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" noChangeAspect="1"/></wp:cNvGraphicFramePr><a:graphic xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"><a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture"><pic:pic xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture"><pic:nvPicPr><pic:cNvPr id="1" name="Picture 1"/><pic:cNvPicPr/></pic:nvPicPr><pic:blipFill><a:blip r:embed="rId10"/><a:stretch><a:fillRect/></a:stretch></pic:blipFill><pic:spPr><a:xfrm><a:off x="0" y="0"/><a:ext cx="914400" cy="914400"/></a:xfrm><a:prstGeom prst="rect"><a:avLst/></a:prstGeom></pic:spPr></pic:pic></a:graphicData></a:graphic></wp:inline></w:drawing>')
ANCHOR = DRAWING.replace('<wp:inline distT="0" distB="0" distL="0" distR="0" wp14:anchorId="1F2E3D4C" wp14:editId="5A6B7C8D">', '<wp:anchor distT="0" distB="0" distL="114300" distR="114300" simplePos="0" relativeHeight="251658240" behindDoc="0" locked="0" layoutInCell="1" allowOverlap="1"><wp:simplePos x="0" y="0"/><wp:positionH relativeFrom="column"><wp:posOffset>0</wp:posOffset></wp:positionH><wp:positionV relativeFrom="paragraph"><wp:posOffset>0</wp:posOffset></wp:positionV>').replace('</wp:inline>', '<wp:wrapSquare wrapText="bothSides"/></wp:anchor>').replace('<wp:effectExtent l="0" t="0" r="0" b="0"/>', '<wp:effectExtent l="0" t="0" r="0" b="0"/><wp:wrapSquare wrapText="bothSides"/>', 0)
ALT = ('<mc:AlternateContent><mc:Choice Requires="wps"><w:drawing><wp:inline distT="0" distB="0" distL="0" distR="0"><wp:extent cx="1828800" cy="457200"/><wp:docPr id="2" name="Text Box 2"/><a:graphic xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"><a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape"><wps:wsp><wps:cNvSpPr txBox="1"/><wps:spPr><a:xfrm><a:off x="0" y="0"/><a:ext cx="1828800" cy="457200"/></a:xfrm><a:prstGeom prst="rect"><a:avLst/></a:prstGeom></wps:spPr><wps:txbx><w:txbxContent><w:p><w:r><w:t>Text inside a text box</w:t></w:r></w:p></w:txbxContent></wps:txbx><wps:bodyPr rot="0" vert="horz" wrap="square" anchor="t" anchorCtr="0"><a:noAutofit/></wps:bodyPr></wps:wsp></a:graphicData></a:graphic></wp:inline></w:drawing></mc:Choice><mc:Fallback><w:pict><v:shapetype id="_x0000_t202" coordsize="21600,21600" o:spt="202" path="m,l,21600r21600,l21600,xe"><v:stroke joinstyle="miter"/><v:path gradientshapeok="t" o:connecttype="rect"/></v:shapetype><v:shape id="Text Box 2" o:spid="_x0000_s1026" type="#_x0000_t202" style="width:2in;height:36pt;visibility:visible"><v:textbox><w:txbxContent><w:p><w:r><w:t>Text inside a text box</w:t></w:r></w:p></w:txbxContent></v:textbox></v:shape></w:pict></mc:Fallback></mc:AlternateContent>')

package("word-images-shapes.docx",
        p("Images and shapes", '<w:pStyle w:val="Heading1"/>')
        + f'<w:p><w:r><w:t xml:space="preserve">Inline picture: </w:t></w:r><w:r><w:rPr><w:noProof/></w:rPr>{DRAWING}</w:r><w:r><w:t xml:space="preserve"> then text.</w:t></w:r></w:p>'
        + f'<w:p><w:r><w:rPr><w:noProof/></w:rPr>{ANCHOR}</w:r><w:r><w:t>Floating picture anchored to this paragraph, text wraps around it. {LOREM}</w:t></w:r></w:p>'
        + f'<w:p><w:r>{ALT}</w:r></w:p>'
        + '<w:p><w:r><w:object w:dxaOrig="1440" w:dyaOrig="1440"><v:shape id="_x0000_i1025" type="#_x0000_t75" style="width:1in;height:1in"><v:imagedata r:id="rId10" o:title=""/></v:shape><o:OLEObject Type="Embed" ProgID="Package" ShapeID="_x0000_i1025" DrawAspect="Content" ObjectID="_1234567890" r:id="rId10"/></w:object></w:r></w:p>',
        extra={**WORD_EXTRA, "word/media/image1.png": ("png", png())},
        rels=WORD_RELS + [("rId10", "image", "media/image1.png", False)])

package("word-fields-toc.docx",
        p("Contents", '<w:pStyle w:val="Heading1"/>')
        + '<w:sdt><w:sdtPr><w:docPartObj><w:docPartGallery w:val="Table of Contents"/><w:docPartUnique/></w:docPartObj></w:sdtPr><w:sdtEndPr><w:rPr><w:b/><w:bCs/></w:rPr></w:sdtEndPr><w:sdtContent><w:p><w:pPr><w:pStyle w:val="TOC1"/><w:tabs><w:tab w:val="right" w:leader="dot" w:pos="9350"/></w:tabs></w:pPr><w:r><w:fldChar w:fldCharType="begin"/></w:r><w:r><w:instrText xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText></w:r><w:r><w:fldChar w:fldCharType="separate"/></w:r><w:hyperlink w:anchor="_Toc1" w:history="1"><w:r><w:rPr><w:rStyle w:val="Hyperlink"/></w:rPr><w:t>Contents</w:t></w:r><w:r><w:tab/></w:r><w:r><w:fldChar w:fldCharType="begin"/></w:r><w:r><w:instrText xml:space="preserve"> PAGEREF _Toc1 \\h </w:instrText></w:r><w:r><w:fldChar w:fldCharType="separate"/></w:r><w:r><w:t>1</w:t></w:r><w:r><w:fldChar w:fldCharType="end"/></w:r></w:hyperlink></w:p><w:p><w:r><w:fldChar w:fldCharType="end"/></w:r></w:p></w:sdtContent></w:sdt>'
        + '<w:p><w:pPr><w:pStyle w:val="Heading1"/></w:pPr><w:bookmarkStart w:id="0" w:name="_Toc1"/><w:r><w:t>Chapter</w:t></w:r><w:bookmarkEnd w:id="0"/></w:p>'
        + '<w:p><w:r><w:t xml:space="preserve">Date: </w:t></w:r><w:fldSimple w:instr=" DATE \\@ &quot;d MMMM yyyy&quot; "><w:r><w:t>1 January 2025</w:t></w:r></w:fldSimple><w:r><w:t xml:space="preserve">. Author: </w:t></w:r><w:fldSimple w:instr=" AUTHOR "><w:r><w:rPr><w:noProof/></w:rPr><w:t>Author</w:t></w:r></w:fldSimple><w:r><w:t>.</w:t></w:r></w:p>'
        + '<w:p><w:r><w:t xml:space="preserve">Nested fields: </w:t></w:r><w:r><w:fldChar w:fldCharType="begin"/></w:r><w:r><w:instrText xml:space="preserve"> IF </w:instrText></w:r><w:r><w:fldChar w:fldCharType="begin"/></w:r><w:r><w:instrText xml:space="preserve"> PAGE </w:instrText></w:r><w:r><w:fldChar w:fldCharType="end"/></w:r><w:r><w:instrText xml:space="preserve"> = 1 "first" "other" </w:instrText></w:r><w:r><w:fldChar w:fldCharType="separate"/></w:r><w:r><w:t>first</w:t></w:r><w:r><w:fldChar w:fldCharType="end"/></w:r></w:p>'
        + '<w:p><w:r><w:t xml:space="preserve">Cross-ref: see </w:t></w:r><w:r><w:fldChar w:fldCharType="begin"/></w:r><w:r><w:instrText xml:space="preserve"> REF _Toc1 \\h </w:instrText></w:r><w:r><w:fldChar w:fldCharType="separate"/></w:r><w:r><w:t>Chapter</w:t></w:r><w:r><w:fldChar w:fldCharType="end"/></w:r><w:r><w:t xml:space="preserve"> and a hyperlink </w:t></w:r><w:hyperlink r:id="rId11" w:tooltip="Example" w:history="1"><w:r><w:rPr><w:rStyle w:val="Hyperlink"/></w:rPr><w:t>https://example.com/a?b=1&amp;c=2</w:t></w:r></w:hyperlink><w:r><w:t>.</w:t></w:r></w:p>',
        extra=WORD_EXTRA, rels=WORD_RELS + [("rId11", "hyperlink", "https://example.com/a?b=1&amp;c=2", True)])

def tc(text, extra=""):
    return f'<w:tc><w:tcPr><w:tcW w:w="2000" w:type="dxa"/>{extra}</w:tcPr><w:p><w:r><w:t>{text}</w:t></w:r></w:p></w:tc>'

package("word-tables.docx",
        p("Tables", '<w:pStyle w:val="Heading1"/>')
        + '<w:tbl><w:tblPr><w:tblStyle w:val="TableGrid"/><w:tblW w:w="0" w:type="auto"/><w:tblLook w:val="04A0" w:firstRow="1" w:lastRow="0" w:firstColumn="1" w:lastColumn="0" w:noHBand="0" w:noVBand="1"/></w:tblPr><w:tblGrid><w:gridCol w:w="2000"/><w:gridCol w:w="2000"/><w:gridCol w:w="2000"/></w:tblGrid>'
        + '<w:tr w:rsidR="00C62A19" w14:paraId="0A0B0C0D"><w:trPr><w:tblHeader/></w:trPr>' + tc("Header 1") + tc("Header 2") + tc("Header 3") + '</w:tr>'
        + '<w:tr>' + tc("Merged", '<w:gridSpan w:val="2"/><w:shd w:val="clear" w:color="auto" w:fill="D9E2F3"/>') + tc("c") + '</w:tr>'
        + '<w:tr>' + tc("vMerge start", '<w:vMerge w:val="restart"/>') + tc("x") + tc("y") + '</w:tr>'
        + '<w:tr>' + tc("", '<w:vMerge/>') + '<w:tc><w:tcPr><w:tcW w:w="2000" w:type="dxa"/></w:tcPr><w:tbl><w:tblPr><w:tblW w:w="0" w:type="auto"/></w:tblPr><w:tblGrid><w:gridCol w:w="900"/></w:tblGrid><w:tr><w:tc><w:p><w:r><w:t>nested</w:t></w:r></w:p></w:tc></w:tr></w:tbl><w:p/></w:tc>' + tc("z", '<w:vAlign w:val="center"/>') + '</w:tr>'
        + '</w:tbl>'
        + p("Between tables.")
        + '<w:tbl><w:tblPr><w:tblW w:w="5000" w:type="pct"/><w:jc w:val="center"/><w:tblBorders><w:top w:val="double" w:sz="6" w:space="0" w:color="FF0000"/></w:tblBorders><w:tblCellMar><w:left w:w="200" w:type="dxa"/></w:tblCellMar></w:tblPr><w:tblGrid><w:gridCol w:w="4000"/></w:tblGrid><w:tr><w:tc><w:p><w:pPr><w:jc w:val="right"/></w:pPr><w:r><w:rPr><w:b/></w:rPr><w:t>Bold right-aligned</w:t></w:r></w:p><w:p><w:r><w:t>Second paragraph in the cell</w:t></w:r></w:p></w:tc></w:tr></w:tbl>'
        + '<w:tbl><w:tblPr><w:tblW w:w="0" w:type="auto"/></w:tblPr><w:tblGrid><w:gridCol w:w="4000"/></w:tblGrid><w:tr><w:tc><w:p><w:r><w:t>Table immediately after another; document ends with a table</w:t></w:r></w:p></w:tc></w:tr></w:tbl>'
        + '<w:p/>',
        extra=WORD_EXTRA, rels=WORD_RELS)

package("word-content-controls.docx",
        '<w:sdt><w:sdtPr><w:alias w:val="Title"/><w:tag w:val="title"/><w:id w:val="-1"/><w:placeholder><w:docPart w:val="DefaultPlaceholder"/></w:placeholder><w:dataBinding w:prefixMappings="xmlns:ns0=\'http://purl.org/dc/elements/1.1/\'" w:xpath="/ns0:title" w:storeItemID="{ABC}"/><w:text/></w:sdtPr><w:sdtContent><w:p><w:pPr><w:pStyle w:val="Heading1"/></w:pPr><w:r><w:t>Bound title</w:t></w:r></w:p></w:sdtContent></w:sdt>'
        + '<w:p><w:r><w:t xml:space="preserve">Inline: </w:t></w:r><w:sdt><w:sdtPr><w:id w:val="2"/><w:date w:fullDate="2025-01-01T00:00:00Z"><w:dateFormat w:val="yyyy-MM-dd"/></w:date></w:sdtPr><w:sdtContent><w:r><w:t>2025-01-01</w:t></w:r></w:sdtContent></w:sdt><w:r><w:t xml:space="preserve"> and a checkbox </w:t></w:r><w:sdt><w:sdtPr><w:id w:val="3"/><w14:checkbox><w14:checked w14:val="1"/><w14:checkedState w14:val="2612" w14:font="MS Gothic"/><w14:uncheckedState w14:val="2610" w14:font="MS Gothic"/></w14:checkbox></w:sdtPr><w:sdtContent><w:r><w:rPr><w:rFonts w:ascii="MS Gothic" w:hAnsi="MS Gothic"/></w:rPr><w:t>☒</w:t></w:r></w:sdtContent></w:sdt><w:r><w:t xml:space="preserve"> and a dropdown </w:t></w:r><w:sdt><w:sdtPr><w:id w:val="4"/><w:dropDownList><w:listItem w:displayText="Red" w:value="r"/><w:listItem w:displayText="Blue" w:value="b"/></w:dropDownList></w:sdtPr><w:sdtContent><w:r><w:t>Red</w:t></w:r></w:sdtContent></w:sdt><w:r><w:t>.</w:t></w:r></w:p>'
        + '<w:sdt><w:sdtPr><w:id w:val="5"/><w:group/></w:sdtPr><w:sdtContent><w:p><w:r><w:t>Group control paragraph one</w:t></w:r></w:p><w:tbl><w:tblPr><w:tblW w:w="0" w:type="auto"/></w:tblPr><w:tblGrid><w:gridCol w:w="4000"/></w:tblGrid><w:tr><w:tc><w:p><w:r><w:t>table in group</w:t></w:r></w:p></w:tc></w:tr></w:tbl><w:p><w:r><w:t>Group control paragraph two</w:t></w:r></w:p></w:sdtContent></w:sdt>'
        + '<w:p><w:sdt><w:sdtPr><w:id w:val="6"/><w:richText/></w:sdtPr><w:sdtContent><w:r><w:t xml:space="preserve">Rich text control with </w:t></w:r><w:r><w:rPr><w:b/></w:rPr><w:t>bold</w:t></w:r><w:sdt><w:sdtPr><w:id w:val="7"/></w:sdtPr><w:sdtContent><w:r><w:t xml:space="preserve"> nested control</w:t></w:r></w:sdtContent></w:sdt></w:sdtContent></w:sdt></w:p>',
        extra=WORD_EXTRA, rels=WORD_RELS)

package("word-symbols-special.docx",
        '<w:p><w:r><w:t xml:space="preserve">Symbol: </w:t></w:r><w:r><w:sym w:font="Wingdings" w:char="F0FC"/></w:r><w:r><w:t xml:space="preserve"> soft hyphen: super</w:t><w:softHyphen/><w:t>califragilistic no-break: 555</w:t><w:noBreakHyphen/><w:t xml:space="preserve">1234 cr:</w:t><w:cr/><w:t>after cr, positional tab:</w:t><w:ptab w:relativeTo="margin" w:alignment="right" w:leader="dot"/><w:t>right</w:t></w:r></w:p>'
        + '<w:p><w:r><w:t xml:space="preserve">Ruby: </w:t></w:r><w:r><w:ruby><w:rubyPr><w:rubyAlign w:val="center"/><w:hps w:val="10"/><w:hpsRaise w:val="18"/><w:hpsBaseText w:val="22"/><w:lid w:val="ja-JP"/></w:rubyPr><w:rt><w:r><w:rPr><w:sz w:val="10"/></w:rPr><w:t>かんじ</w:t></w:r></w:rt><w:rubyBase><w:r><w:t>漢字</w:t></w:r></w:rubyBase></w:ruby></w:r><w:r><w:t xml:space="preserve"> and math </w:t></w:r><m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"><m:sSup><m:e><m:r><m:t>x</m:t></m:r></m:e><m:sup><m:r><m:t>2</m:t></m:r></m:sup></m:sSup></m:oMath><w:r><w:t>.</w:t></w:r></w:p>'
        + '<m:oMathPara xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"><m:oMath><m:f><m:num><m:r><m:t>a</m:t></m:r></m:num><m:den><m:r><m:t>b</m:t></m:r></m:den></m:f></m:oMath></m:oMathPara>'
        + '<w:p><w:r><w:rPr><w:rFonts w:ascii="Consolas" w:hAnsi="Consolas"/><w:spacing w:val="20"/><w:w w:val="150"/><w:kern w:val="28"/><w:position w:val="6"/><w:effect w:val="none"/><w:bdr w:val="single" w:sz="4" w:space="0" w:color="auto"/><w:shd w:val="clear" w:color="auto" w:fill="FFFF00"/><w:fitText w:val="2000" w:id="1"/><w:em w:val="dot"/><w:lang w:val="fr-FR" w:eastAsia="ja-JP" w:bidi="he-IL"/><w:eastAsianLayout w:id="1" w:combine="1"/><w:specVanish/></w:rPr><w:t>Every exotic run property</w:t></w:r><w:r><w:rPr><w:outline/><w:shadow/><w:emboss/><w:imprint/><w:vanish/><w:webHidden/><w:snapToGrid w:val="0"/><w:noProof/><w:rtl/><w:cs/><w:color w:val="auto" w:themeColor="text1"/><w:u w:val="wavyDouble" w:color="FF0000"/><w:highlight w:val="lightGray"/><w14:glow w14:rad="63500"><w14:srgbClr w14:val="FFC000"/></w14:glow><w14:shadow w14:blurRad="50800" w14:dist="38100" w14:dir="2700000" w14:sx="100000" w14:sy="100000" w14:kx="0" w14:ky="0" w14:algn="tl"><w14:srgbClr w14:val="000000"/></w14:shadow><w14:ligatures w14:val="standardContextual"/></w:rPr><w:t xml:space="preserve"> and more</w:t></w:r></w:p>'
        + '<w:p><w:pPr><w:framePr w:w="3000" w:hAnchor="page" w:x="1000" w:y="1000"/><w:pBdr><w:top w:val="single" w:sz="4" w:space="1" w:color="auto"/><w:left w:val="dashed" w:sz="8" w:space="4" w:color="FF0000"/><w:bottom w:val="double" w:sz="6" w:space="1" w:color="auto"/><w:right w:val="dotted" w:sz="4" w:space="4" w:color="00FF00"/><w:between w:val="single" w:sz="4" w:space="1" w:color="auto"/><w:bar w:val="thick" w:sz="12" w:space="0" w:color="auto"/></w:pBdr><w:shd w:val="pct25" w:color="auto" w:fill="E7E6E6" w:themeFill="background2"/><w:suppressAutoHyphens/><w:kinsoku w:val="0"/><w:wordWrap w:val="0"/><w:overflowPunct w:val="0"/><w:topLinePunct/><w:autoSpaceDE w:val="0"/><w:autoSpaceDN w:val="0"/><w:bidi/><w:adjustRightInd w:val="0"/><w:snapToGrid w:val="0"/><w:spacing w:before="120" w:beforeAutospacing="1" w:after="120" w:afterAutospacing="1" w:line="300" w:lineRule="exact"/><w:ind w:start="720" w:end="360" w:firstLineChars="200" w:firstLine="480"/><w:contextualSpacing/><w:mirrorIndents/><w:suppressOverlap/><w:jc w:val="distribute"/><w:textDirection w:val="btLr"/><w:textAlignment w:val="center"/><w:textboxTightWrap w:val="allLines"/><w:outlineLvl w:val="3"/><w:divId w:val="123"/><w:cnfStyle w:val="100000000000" w:firstRow="1"/><w:rPr><w:b/><w:sz w:val="28"/></w:rPr></w:pPr><w:r><w:t>Every exotic paragraph property</w:t></w:r></w:p>'
        + '<w:p><w:pPr><w:pStyle w:val="Normal"/><w:widowControl w:val="0"/><w:tabs><w:tab w:val="clear" w:pos="720"/><w:tab w:val="left" w:pos="1440"/><w:tab w:val="center" w:leader="hyphen" w:pos="4320"/><w:tab w:val="right" w:leader="underscore" w:pos="8640"/><w:tab w:val="decimal" w:leader="middleDot" w:pos="6000"/><w:tab w:val="bar" w:pos="7000"/><w:tab w:val="num" w:pos="360"/></w:tabs><w:suppressLineNumbers/><w:keepNext w:val="0"/><w:keepLines w:val="false"/><w:pageBreakBefore w:val="on"/></w:pPr><w:r><w:t>Tabs and toggles</w:t></w:r></w:p>',
        extra=WORD_EXTRA, rels=WORD_RELS)

package("word-bookmarks.docx",
        '<w:bookmarkStart w:id="0" w:name="_GoBack"/><w:bookmarkEnd w:id="0"/>'
        + '<w:p><w:bookmarkStart w:id="1" w:name="top"/><w:bookmarkEnd w:id="1"/><w:r><w:t xml:space="preserve">Text with </w:t></w:r><w:bookmarkStart w:id="2" w:name="inner"/><w:r><w:t>a bookmarked span</w:t></w:r><w:bookmarkEnd w:id="2"/><w:r><w:t>.</w:t></w:r></w:p>'
        + '<w:p><w:bookmarkStart w:id="3" w:name="spanning"/><w:r><w:t>Bookmark spans paragraphs</w:t></w:r></w:p>'
        + '<w:p><w:r><w:t>and ends here</w:t></w:r><w:bookmarkEnd w:id="3"/></w:p>'
        + '<w:bookmarkStart w:id="4" w:name="before_table" w:colFirst="0" w:colLast="1"/><w:tbl><w:tblPr><w:tblW w:w="0" w:type="auto"/></w:tblPr><w:tblGrid><w:gridCol w:w="2000"/><w:gridCol w:w="2000"/></w:tblGrid><w:tr><w:tc><w:p><w:r><w:t>a</w:t></w:r></w:p></w:tc><w:tc><w:p><w:r><w:t>b</w:t></w:r></w:p></w:tc></w:tr></w:tbl><w:bookmarkEnd w:id="4"/>'
        + '<w:p><w:r><w:t>Last paragraph</w:t></w:r></w:p><w:bookmarkStart w:id="5" w:name="end_of_doc"/><w:bookmarkEnd w:id="5"/>',
        extra=WORD_EXTRA, rels=WORD_RELS)

package("word-long-rendered.docx",
        "".join(
            f'<w:p w:rsidR="00C62A19" w:rsidRDefault="00C62A19" w14:paraId="{i:08X}" w14:textId="{i * 7:08X}"><w:pPr><w:pStyle w:val="{"Heading2" if i % 9 == 0 else "Normal"}"/></w:pPr>'
            + (f'<w:r><w:lastRenderedPageBreak/><w:t xml:space="preserve">' if i % 8 == 0 and i else '<w:r><w:t xml:space="preserve">')
            + f"Paragraph {i}. {LOREM if i % 9 else ''}</w:t></w:r></w:p>"
            for i in range(120)
        ),
        extra=WORD_EXTRA, rels=WORD_RELS)

package("word-styles-many.docx",
        "".join(p(f"Styled {s}", f'<w:pStyle w:val="{s}"/>') for s in ["Normal", "Heading1", "Heading2", "ListParagraph", "FootnoteText", "Missing Style Id", "Title"]),
        styles=STYLES_WORD.replace("</w:styles>", '<w:style w:type="paragraph" w:styleId="Title"><w:name w:val="Title"/><w:basedOn w:val="Normal"/><w:next w:val="Normal"/><w:link w:val="TitleChar"/><w:qFormat/><w:pPr><w:spacing w:after="0" w:line="240" w:lineRule="auto"/><w:contextualSpacing/></w:pPr><w:rPr><w:rFonts w:asciiTheme="majorHAnsi" w:hAnsiTheme="majorHAnsi"/><w:spacing w:val="-10"/><w:kern w:val="28"/><w:sz w:val="56"/></w:rPr></w:style><w:style w:type="paragraph" w:customStyle="1" w:styleId="MyBody"><w:name w:val="My Body"/><w:basedOn w:val="Normal"/><w:pPr><w:ind w:left="1440" w:hanging="720"/><w:jc w:val="both"/></w:pPr><w:rPr><w:rFonts w:ascii="Georgia" w:hAnsi="Georgia"/><w:sz w:val="24"/></w:rPr></w:style><w:style w:type="paragraph" w:customStyle="1" w:styleId="Cyclic1"><w:name w:val="Cyclic 1"/><w:basedOn w:val="Cyclic2"/></w:style><w:style w:type="paragraph" w:customStyle="1" w:styleId="Cyclic2"><w:name w:val="Cyclic 2"/><w:basedOn w:val="Cyclic1"/></w:style></w:styles>')
        + "", extra=WORD_EXTRA, rels=WORD_RELS)
# (the body above references MyBody/Cyclic through a second doc)
package("word-style-chain.docx",
        p("Custom body style", '<w:pStyle w:val="MyBody"/>') + p("Cyclic style must not hang", '<w:pStyle w:val="Cyclic1"/>') + p("Direct overrides", '<w:pStyle w:val="MyBody"/><w:ind w:left="0" w:firstLine="0"/><w:jc w:val="left"/>', '<w:rFonts w:ascii="Arial" w:hAnsi="Arial"/><w:sz w:val="18"/>'),
        styles=STYLES_WORD.replace("</w:styles>", '<w:style w:type="paragraph" w:customStyle="1" w:styleId="MyBody"><w:name w:val="My Body"/><w:basedOn w:val="Normal"/><w:pPr><w:ind w:left="1440" w:hanging="720"/><w:jc w:val="both"/></w:pPr><w:rPr><w:rFonts w:ascii="Georgia" w:hAnsi="Georgia"/><w:sz w:val="24"/></w:rPr></w:style><w:style w:type="paragraph" w:customStyle="1" w:styleId="Cyclic1"><w:name w:val="Cyclic 1"/><w:basedOn w:val="Cyclic2"/></w:style><w:style w:type="paragraph" w:customStyle="1" w:styleId="Cyclic2"><w:name w:val="Cyclic 2"/><w:basedOn w:val="Cyclic1"/></w:style></w:styles>'),
        extra=WORD_EXTRA, rels=WORD_RELS)

package("word-background-glossary.docx",
        p("Document with a page background and extra parts."),
        ns=NS_WORD, extra={**WORD_EXTRA, "word/glossary/document.xml": ("glossary", f'{DECL}<w:glossaryDocument {NS_MIN}><w:docParts><w:docPart><w:docPartPr><w:name w:val="DefaultPlaceholder"/></w:docPartPr><w:docPartBody><w:p><w:r><w:t>Click here.</w:t></w:r></w:p></w:docPartBody></w:docPart></w:docParts></w:glossaryDocument>'), "customXml/item1.xml": ("custom", '<?xml version="1.0" encoding="UTF-8" standalone="no"?><b:Sources xmlns:b="http://schemas.openxmlformats.org/officeDocument/2006/bibliography"/>')},
        rels=WORD_RELS)
# patch in a w:background between <w:document> and <w:body>
import re as _re
with zipfile.ZipFile(f"{out_dir}/word-background-glossary.docx") as z:
    parts = {n: z.read(n) for n in z.namelist()}
parts["word/document.xml"] = parts["word/document.xml"].replace(b"<w:body>", b'<w:background w:color="FFFFCC"><v:background id="_x0000_s1025" o:bwmode="white" fillcolor="#ffc"><v:fill r:id="rId10" o:title="" type="tile"/></v:background></w:background><w:body>', 1)
save("word-background-glossary.docx", parts)

# --- Google Docs style ------------------------------------------------------
NS_G = 'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture" xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math" xmlns:v="urn:schemas-microsoft-com:vml" xmlns:o="urn:schemas-microsoft-com:office:office" xmlns:w10="urn:schemas-microsoft-com:office:word"'
STYLES_G = f'{DECL}<w:styles {NS_MIN}><w:docDefaults><w:rPrDefault><w:rPr><w:rFonts w:ascii="Arial" w:cs="Arial" w:eastAsia="Arial" w:hAnsi="Arial"/><w:sz w:val="22"/><w:szCs w:val="22"/><w:lang w:val="en"/></w:rPr></w:rPrDefault><w:pPrDefault><w:pPr><w:spacing w:after="0" w:before="0" w:line="276" w:lineRule="auto"/></w:pPr></w:pPrDefault></w:docDefaults><w:style w:type="paragraph" w:styleId="Normal" w:default="1"><w:name w:val="normal"/></w:style><w:style w:type="table" w:styleId="TableNormal" w:default="1"><w:name w:val="Table Normal"/><w:tblPr><w:tblCellMar><w:top w:w="0.0" w:type="dxa"/><w:left w:w="0.0" w:type="dxa"/><w:bottom w:w="0.0" w:type="dxa"/><w:right w:w="0.0" w:type="dxa"/></w:tblCellMar></w:tblPr></w:style><w:style w:type="paragraph" w:styleId="Heading1"><w:name w:val="heading 1"/><w:basedOn w:val="Normal"/><w:next w:val="Normal"/><w:pPr><w:keepNext w:val="1"/><w:keepLines w:val="1"/><w:spacing w:after="120" w:before="480" w:lineRule="auto"/><w:outlineLvl w:val="0"/></w:pPr><w:rPr><w:b w:val="1"/><w:sz w:val="48"/><w:szCs w:val="48"/></w:rPr></w:style><w:style w:type="paragraph" w:styleId="Title"><w:name w:val="Title"/><w:basedOn w:val="Normal"/><w:next w:val="Normal"/><w:pPr><w:keepNext w:val="1"/><w:keepLines w:val="1"/><w:spacing w:after="60" w:before="0" w:lineRule="auto"/></w:pPr><w:rPr><w:sz w:val="52"/><w:szCs w:val="52"/></w:rPr></w:style><w:style w:type="paragraph" w:styleId="Subtitle"><w:name w:val="Subtitle"/><w:basedOn w:val="Normal"/><w:next w:val="Normal"/><w:pPr><w:keepNext w:val="1"/><w:keepLines w:val="1"/><w:spacing w:after="320" w:before="0" w:lineRule="auto"/></w:pPr><w:rPr><w:rFonts w:ascii="Georgia" w:cs="Georgia" w:eastAsia="Georgia" w:hAnsi="Georgia"/><w:i w:val="1"/><w:color w:val="666666"/><w:sz w:val="48"/><w:szCs w:val="48"/></w:rPr></w:style></w:styles>'
NUMBERING_G = f'{DECL}<w:numbering {NS_MIN}><w:abstractNum w:abstractNumId="0"><w:lvl w:ilvl="0"><w:start w:val="1"/><w:numFmt w:val="bullet"/><w:lvlText w:val="●"/><w:lvlJc w:val="left"/><w:pPr><w:ind w:left="720" w:hanging="360"/></w:pPr><w:rPr><w:u w:val="none"/><w:vertAlign w:val="baseline"/></w:rPr></w:lvl><w:lvl w:ilvl="1"><w:start w:val="1"/><w:numFmt w:val="bullet"/><w:lvlText w:val="○"/><w:lvlJc w:val="left"/><w:pPr><w:ind w:left="1440" w:hanging="360"/></w:pPr><w:rPr><w:u w:val="none"/><w:vertAlign w:val="baseline"/></w:rPr></w:lvl><w:lvl w:ilvl="2"><w:start w:val="1"/><w:numFmt w:val="bullet"/><w:lvlText w:val="■"/><w:lvlJc w:val="left"/><w:pPr><w:ind w:left="2160" w:hanging="360"/></w:pPr><w:rPr><w:u w:val="none"/><w:vertAlign w:val="baseline"/></w:rPr></w:lvl></w:abstractNum><w:abstractNum w:abstractNumId="1"><w:lvl w:ilvl="0"><w:start w:val="1"/><w:numFmt w:val="decimal"/><w:lvlText w:val="%1."/><w:lvlJc w:val="left"/><w:pPr><w:ind w:left="720" w:hanging="360"/></w:pPr><w:rPr><w:u w:val="none"/><w:vertAlign w:val="baseline"/></w:rPr></w:lvl><w:lvl w:ilvl="1"><w:start w:val="1"/><w:numFmt w:val="lowerLetter"/><w:lvlText w:val="%2."/><w:lvlJc w:val="left"/><w:pPr><w:ind w:left="1440" w:hanging="360"/></w:pPr><w:rPr><w:u w:val="none"/><w:vertAlign w:val="baseline"/></w:rPr></w:lvl></w:abstractNum><w:num w:numId="1"><w:abstractNumId w:val="0"/></w:num><w:num w:numId="2"><w:abstractNumId w:val="1"/></w:num></w:numbering>'
G_SECT = '<w:sectPr><w:pgSz w:h="15840" w:w="12240" w:orient="portrait"/><w:pgMar w:bottom="1440" w:top="1440" w:left="1440" w:right="1440" w:header="720" w:footer="720"/><w:pgNumType w:start="1"/></w:sectPr>'

def gp(text, ppr="", rpr=""):
    pp = f"<w:pPr>{ppr}<w:rPr>{rpr}</w:rPr></w:pPr>"
    rp = f"<w:rPr>{rpr}</w:rPr>" if rpr else "<w:rPr/>"
    if not text:
        return f'<w:p w:rsidR="00000000" w:rsidDel="00000000" w:rsidP="00000000" w:rsidRDefault="00000000" w:rsidRPr="00000000" w14:paraId="00000000" xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml">{pp}</w:p>'
    return f'<w:p w:rsidR="00000000" w:rsidDel="00000000" w:rsidP="00000000" w:rsidRDefault="00000000" w:rsidRPr="00000000" w14:paraId="00000000" xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml">{pp}<w:r w:rsidDel="00000000" w:rsidR="00000000" w:rsidRPr="00000000">{rp}<w:t xml:space="preserve">{text}</w:t></w:r></w:p>'

package("gdocs-basic.docx",
        gp("Google Docs export", '<w:pStyle w:val="Title"/>', "") + gp("A subtitle", '<w:pStyle w:val="Subtitle"/>') + gp("Heading one", '<w:pStyle w:val="Heading1"/>')
        + gp("Body text in Arial with explicit rtl off.", '<w:rPr><w:rtl w:val="0"/></w:rPr>', '<w:rtl w:val="0"/>')
        + gp("Bulleted item", '<w:numPr><w:ilvl w:val="0"/><w:numId w:val="1"/></w:numPr><w:ind w:left="720" w:hanging="360"/>')
        + gp("Sub bullet", '<w:numPr><w:ilvl w:val="1"/><w:numId w:val="1"/></w:numPr><w:ind w:left="1440" w:hanging="360"/>')
        + gp("Numbered", '<w:numPr><w:ilvl w:val="0"/><w:numId w:val="2"/></w:numPr><w:ind w:left="720" w:hanging="360"/>')
        + gp("", "") + gp("Centered", '<w:jc w:val="center"/>', '<w:b w:val="1"/><w:sz w:val="28"/><w:szCs w:val="28"/>'),
        styles=STYLES_G, ns=NS_G, decl='<?xml version="1.0" encoding="UTF-8" standalone="yes"?>', sect=G_SECT,
        extra={"word/numbering.xml": ("numbering", NUMBERING_G), "word/settings.xml": ("settings", f'{DECL}<w:settings {NS_MIN}><w:displayBackgroundShape w:val="1"/><w:defaultTabStop w:val="720"/><w:compat><w:compatSetting w:name="compatibilityMode" w:uri="http://schemas.microsoft.com/office/word" w:val="15"/></w:compat></w:settings>')},
        rels=[("rId2", "numbering", "numbering.xml", False), ("rId3", "settings", "settings.xml", False)],
        order=["_rels/.rels", "word/document.xml", "word/styles.xml", "word/numbering.xml", "word/settings.xml", "word/_rels/document.xml.rels", "[Content_Types].xml"])

package("gdocs-table-image.docx",
        gp("Table from Google Docs")
        + '<w:tbl><w:tblPr><w:tblStyle w:val="Table1"/><w:tblW w:w="9360.0" w:type="dxa"/><w:jc w:val="left"/><w:tblBorders><w:top w:color="000000" w:space="0" w:sz="4" w:val="single"/><w:left w:color="000000" w:space="0" w:sz="4" w:val="single"/><w:bottom w:color="000000" w:space="0" w:sz="4" w:val="single"/><w:right w:color="000000" w:space="0" w:sz="4" w:val="single"/><w:insideH w:color="000000" w:space="0" w:sz="4" w:val="single"/><w:insideV w:color="000000" w:space="0" w:sz="4" w:val="single"/></w:tblBorders><w:tblLayout w:type="fixed"/><w:tblLook w:val="0000"/></w:tblPr><w:tblGrid><w:gridCol w:w="4680"/><w:gridCol w:w="4680"/><w:tblGridChange w:id="0"><w:tblGrid><w:gridCol w:w="4680"/><w:gridCol w:w="4680"/></w:tblGrid></w:tblGridChange></w:tblGrid><w:tr><w:trPr><w:cantSplit w:val="0"/><w:tblHeader w:val="0"/></w:trPr><w:tc><w:tcPr><w:shd w:fill="auto" w:val="clear"/><w:tcMar><w:top w:w="100.0" w:type="dxa"/><w:left w:w="100.0" w:type="dxa"/><w:bottom w:w="100.0" w:type="dxa"/><w:right w:w="100.0" w:type="dxa"/></w:tcMar></w:tcPr>' + gp("A") + '</w:tc><w:tc><w:tcPr><w:shd w:fill="auto" w:val="clear"/></w:tcPr>' + gp("B") + '</w:tc></w:tr></w:tbl>'
        + gp("").replace("</w:pPr></w:p>", f"</w:pPr><w:r><w:rPr/>{DRAWING.replace('rId10', 'rId4')}</w:r></w:p>"),
        styles=STYLES_G, ns=NS_G, sect=G_SECT,
        extra={"word/media/image1.png": ("png", png(3, 3, (0, 0, 0)))},
        rels=[("rId4", "image", "media/image1.png", False)])

# --- LibreOffice style ------------------------------------------------------
NS_LO = 'xmlns:o="urn:schemas-microsoft-com:office:office" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" xmlns:v="urn:schemas-microsoft-com:vml" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xmlns:w10="urn:schemas-microsoft-com:office:word" xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape" xmlns:wpg="http://schemas.microsoft.com/office/word/2010/wordprocessingGroup" xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006" xmlns:wp14="http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing" xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml" xmlns:w15="http://schemas.microsoft.com/office/word/2012/wordml" mc:Ignorable="w14 wp14 w15"'
STYLES_LO = f'{DECL}<w:styles {NS_MIN} xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml" xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006" mc:Ignorable="w14"><w:docDefaults><w:rPrDefault><w:rPr><w:rFonts w:ascii="Liberation Serif" w:hAnsi="Liberation Serif" w:eastAsia="NSimSun" w:cs="Lucida Sans"/><w:kern w:val="2"/><w:sz w:val="24"/><w:szCs w:val="24"/><w:lang w:val="en-US" w:eastAsia="zh-CN" w:bidi="hi-IN"/></w:rPr></w:rPrDefault><w:pPrDefault><w:pPr><w:widowControl/></w:pPr></w:pPrDefault></w:docDefaults><w:style w:type="paragraph" w:styleId="Normal"><w:name w:val="Normal"/><w:qFormat/><w:pPr><w:widowControl/><w:bidi w:val="0"/><w:jc w:val="left"/></w:pPr><w:rPr><w:rFonts w:ascii="Liberation Serif" w:hAnsi="Liberation Serif" w:eastAsia="NSimSun" w:cs="Lucida Sans"/><w:color w:val="auto"/><w:kern w:val="2"/><w:sz w:val="24"/><w:szCs w:val="24"/><w:lang w:val="en-US" w:eastAsia="zh-CN" w:bidi="hi-IN"/></w:rPr></w:style><w:style w:type="paragraph" w:styleId="Heading1"><w:name w:val="Heading 1"/><w:basedOn w:val="Heading"/><w:next w:val="TextBody"/><w:qFormat/><w:pPr><w:numPr><w:ilvl w:val="0"/><w:numId w:val="1"/></w:numPr><w:spacing w:before="240" w:after="120"/><w:outlineLvl w:val="0"/></w:pPr><w:rPr><w:b/><w:bCs/><w:sz w:val="36"/><w:szCs w:val="36"/></w:rPr></w:style><w:style w:type="paragraph" w:styleId="Heading"><w:name w:val="Heading"/><w:basedOn w:val="Normal"/><w:next w:val="TextBody"/><w:qFormat/><w:pPr><w:keepNext w:val="true"/><w:spacing w:before="240" w:after="120"/></w:pPr><w:rPr><w:rFonts w:ascii="Liberation Sans" w:hAnsi="Liberation Sans" w:eastAsia="Microsoft YaHei" w:cs="Lucida Sans"/><w:sz w:val="28"/><w:szCs w:val="28"/></w:rPr></w:style><w:style w:type="paragraph" w:styleId="TextBody"><w:name w:val="Body Text"/><w:basedOn w:val="Normal"/><w:pPr><w:spacing w:lineRule="auto" w:line="276" w:before="0" w:after="140"/></w:pPr><w:rPr></w:rPr></w:style><w:style w:type="paragraph" w:styleId="ListParagraph"><w:name w:val="List Paragraph"/><w:basedOn w:val="Normal"/><w:qFormat/><w:pPr><w:spacing w:lineRule="auto" w:line="259" w:before="0" w:after="160"/><w:ind w:left="720" w:hanging="0"/><w:contextualSpacing/></w:pPr><w:rPr></w:rPr></w:style><w:style w:type="character" w:styleId="ListLabel1"><w:name w:val="ListLabel 1"/><w:qFormat/><w:rPr><w:rFonts w:ascii="OpenSymbol;Arial Unicode MS" w:hAnsi="OpenSymbol;Arial Unicode MS" w:cs="OpenSymbol;Arial Unicode MS"/></w:rPr></w:style><w:style w:type="character" w:styleId="InternetLink"><w:name w:val="Hyperlink"/><w:rPr><w:color w:val="000080"/><w:u w:val="single"/></w:rPr></w:style></w:styles>'
NUMBERING_LO = f'{DECL}<w:numbering {NS_MIN}><w:abstractNum w:abstractNumId="1"><w:lvl w:ilvl="0"><w:start w:val="1"/><w:numFmt w:val="none"/><w:suff w:val="nothing"/><w:lvlText w:val=""/><w:lvlJc w:val="left"/><w:pPr><w:tabs><w:tab w:val="num" w:pos="0"/></w:tabs><w:ind w:left="0" w:hanging="0"/></w:pPr></w:lvl></w:abstractNum><w:abstractNum w:abstractNumId="2"><w:lvl w:ilvl="0"><w:start w:val="1"/><w:numFmt w:val="bullet"/><w:lvlText w:val="•"/><w:lvlJc w:val="left"/><w:pPr><w:tabs><w:tab w:val="num" w:pos="720"/></w:tabs><w:ind w:left="720" w:hanging="360"/></w:pPr><w:rPr><w:rStyle w:val="ListLabel1"/><w:rFonts w:ascii="OpenSymbol;Arial Unicode MS" w:hAnsi="OpenSymbol;Arial Unicode MS" w:cs="OpenSymbol;Arial Unicode MS"/></w:rPr></w:lvl><w:lvl w:ilvl="1"><w:start w:val="1"/><w:numFmt w:val="bullet"/><w:lvlText w:val="◦"/><w:lvlJc w:val="left"/><w:pPr><w:tabs><w:tab w:val="num" w:pos="1440"/></w:tabs><w:ind w:left="1440" w:hanging="360"/></w:pPr><w:rPr><w:rStyle w:val="ListLabel1"/></w:rPr></w:lvl></w:abstractNum><w:abstractNum w:abstractNumId="3"><w:lvl w:ilvl="0"><w:start w:val="1"/><w:numFmt w:val="decimal"/><w:lvlText w:val="%1."/><w:lvlJc w:val="left"/><w:pPr><w:tabs><w:tab w:val="num" w:pos="720"/></w:tabs><w:ind w:left="720" w:hanging="360"/></w:pPr></w:lvl><w:lvl w:ilvl="1"><w:start w:val="1"/><w:numFmt w:val="lowerLetter"/><w:lvlText w:val="%2."/><w:lvlJc w:val="left"/><w:pPr><w:tabs><w:tab w:val="num" w:pos="1440"/></w:tabs><w:ind w:left="1440" w:hanging="360"/></w:pPr></w:lvl></w:abstractNum><w:num w:numId="1"><w:abstractNumId w:val="1"/></w:num><w:num w:numId="2"><w:abstractNumId w:val="2"/></w:num><w:num w:numId="3"><w:abstractNumId w:val="3"/></w:num></w:numbering>'
LO_SECT = '<w:sectPr><w:type w:val="nextPage"/><w:pgSz w:w="11906" w:h="16838"/><w:pgMar w:left="1134" w:right="1134" w:gutter="0" w:header="0" w:top="1134" w:footer="0" w:bottom="1134"/><w:pgNumType w:fmt="decimal"/><w:formProt w:val="false"/><w:textDirection w:val="lrTb"/><w:docGrid w:type="default" w:linePitch="100" w:charSpace="0"/></w:sectPr>'

def lp(text, ppr="", rpr=""):
    return f'<w:p><w:pPr><w:pStyle w:val="Normal"/>{ppr}<w:rPr></w:rPr></w:pPr><w:r><w:rPr>{rpr}</w:rPr><w:t>{text}</w:t></w:r></w:p>'

package("lo-basic.docx",
        '<w:p><w:pPr><w:pStyle w:val="Heading1"/><w:rPr></w:rPr></w:pPr><w:r><w:rPr></w:rPr><w:t>LibreOffice Writer export</w:t></w:r></w:p>'
        + '<w:p><w:pPr><w:pStyle w:val="TextBody"/><w:rPr></w:rPr></w:pPr><w:r><w:rPr></w:rPr><w:t xml:space="preserve">Body text with </w:t></w:r><w:r><w:rPr><w:b/><w:bCs/></w:rPr><w:t>bold</w:t></w:r><w:r><w:rPr></w:rPr><w:t xml:space="preserve"> and </w:t></w:r><w:r><w:rPr><w:i/><w:iCs/></w:rPr><w:t>italic</w:t></w:r><w:r><w:rPr></w:rPr><w:t>.</w:t></w:r></w:p>'
        + lp("Bullet", '<w:numPr><w:ilvl w:val="0"/><w:numId w:val="2"/></w:numPr>').replace('w:val="Normal"', 'w:val="ListParagraph"', 1)
        + lp("Sub bullet", '<w:numPr><w:ilvl w:val="1"/><w:numId w:val="2"/></w:numPr>').replace('w:val="Normal"', 'w:val="ListParagraph"', 1)
        + lp("Numbered", '<w:numPr><w:ilvl w:val="0"/><w:numId w:val="3"/></w:numPr>')
        + lp("Right aligned with explicit bidi", '<w:bidi w:val="0"/><w:jc w:val="right"/>')
        + lp("Hyperlink: ").replace("</w:p>", '<w:hyperlink r:id="rId3"><w:r><w:rPr><w:rStyle w:val="InternetLink"/></w:rPr><w:t>https://www.libreoffice.org/</w:t></w:r></w:hyperlink></w:p>')
        + '<w:p><w:pPr><w:pStyle w:val="Normal"/><w:rPr></w:rPr></w:pPr><w:r><w:br w:type="page"/></w:r><w:r><w:rPr></w:rPr><w:t>After a page break run.</w:t></w:r></w:p>'
        + '<w:tbl><w:tblPr><w:tblW w:w="9638" w:type="dxa"/><w:jc w:val="left"/><w:tblInd w:w="0" w:type="dxa"/><w:tblBorders><w:top w:val="single" w:sz="2" w:space="0" w:color="000000"/><w:left w:val="single" w:sz="2" w:space="0" w:color="000000"/><w:bottom w:val="single" w:sz="2" w:space="0" w:color="000000"/><w:insideH w:val="single" w:sz="2" w:space="0" w:color="000000"/></w:tblBorders><w:tblCellMar><w:top w:w="55" w:type="dxa"/><w:left w:w="55" w:type="dxa"/><w:bottom w:w="55" w:type="dxa"/><w:right w:w="55" w:type="dxa"/></w:tblCellMar></w:tblPr><w:tblGrid><w:gridCol w:w="4819"/><w:gridCol w:w="4819"/></w:tblGrid><w:tr><w:trPr></w:trPr><w:tc><w:tcPr><w:tcW w:w="4819" w:type="dxa"/><w:tcBorders><w:top w:val="single" w:sz="2" w:space="0" w:color="000000"/><w:left w:val="single" w:sz="2" w:space="0" w:color="000000"/><w:bottom w:val="single" w:sz="2" w:space="0" w:color="000000"/></w:tcBorders><w:shd w:fill="auto" w:val="clear"/></w:tcPr>' + lp("A1") + '</w:tc><w:tc><w:tcPr><w:tcW w:w="4819" w:type="dxa"/></w:tcPr>' + lp("B1") + '</w:tc></w:tr></w:tbl>'
        + lp(""),
        styles=STYLES_LO, ns=NS_LO, sect=LO_SECT,
        extra={"word/numbering.xml": ("numbering", NUMBERING_LO), "word/settings.xml": ("settings", f'{DECL}<w:settings {NS_MIN}><w:zoom w:percent="100"/><w:defaultTabStop w:val="709"/><w:autoHyphenation w:val="true"/><w:compat><w:compatSetting w:name="compatibilityMode" w:uri="http://schemas.microsoft.com/office/word" w:val="15"/></w:compat><w:themeFontLang w:val="" w:eastAsia="" w:bidi=""/></w:settings>'), "word/fontTable.xml": ("fontTable", f'{DECL}<w:fonts {NS_MIN}><w:font w:name="Liberation Serif"><w:altName w:val="Times New Roman"/><w:charset w:val="00"/><w:family w:val="roman"/><w:pitch w:val="variable"/></w:font><w:font w:name="OpenSymbol"><w:charset w:val="02"/><w:family w:val="auto"/><w:pitch w:val="default"/></w:font></w:fonts>'), "docProps/core.xml": ("core", CORE.replace("Author", "Writer user")), "docProps/app.xml": ("app", APP.replace("Microsoft Office Word", "LibreOffice/24.2.0.3$Linux_X86_64 LibreOffice_project/"))},
        rels=[("rId2", "numbering", "numbering.xml", False), ("rId3", "hyperlink", "https://www.libreoffice.org/", True), ("rId4", "settings", "settings.xml", False), ("rId5", "fontTable", "fontTable.xml", False)],
        stored=("mimetype",))

package("lo-frames-footnotes.docx",
        '<w:p><w:pPr><w:pStyle w:val="Normal"/><w:rPr></w:rPr></w:pPr><w:r><w:rPr></w:rPr><w:t xml:space="preserve">Text with footnote</w:t></w:r><w:r><w:rPr><w:rStyle w:val="FootnoteAnchor"/></w:rPr><w:footnoteReference w:id="2"/></w:r><w:r><w:rPr></w:rPr><w:t>.</w:t></w:r></w:p>'
        + '<w:p><w:pPr><w:pStyle w:val="Normal"/><w:rPr></w:rPr></w:pPr><w:r><w:rPr></w:rPr><mc:AlternateContent><mc:Choice Requires="wps"><w:drawing><wp:anchor behindDoc="0" distT="0" distB="0" distL="0" distR="0" simplePos="0" locked="0" layoutInCell="0" allowOverlap="1" relativeHeight="2"><wp:simplePos x="0" y="0"/><wp:positionH relativeFrom="column"><wp:posOffset>0</wp:posOffset></wp:positionH><wp:positionV relativeFrom="paragraph"><wp:posOffset>0</wp:posOffset></wp:positionV><wp:extent cx="2540000" cy="635000"/><wp:effectExtent l="0" t="0" r="0" b="0"/><wp:wrapSquare wrapText="largest"/><wp:docPr id="1" name="Frame1"/><a:graphic xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"><a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape"><wps:wsp><wps:cNvSpPr txBox="1"/><wps:spPr><a:xfrm><a:off x="0" y="0"/><a:ext cx="2540000" cy="635000"/></a:xfrm><a:prstGeom prst="rect"><a:avLst/></a:prstGeom></wps:spPr><wps:txbx><w:txbxContent><w:p><w:pPr><w:pStyle w:val="Normal"/><w:rPr></w:rPr></w:pPr><w:r><w:rPr></w:rPr><w:t>Frame content</w:t></w:r></w:p></w:txbxContent></wps:txbx><wps:bodyPr anchor="t" lIns="0" tIns="0" rIns="0" bIns="0"><a:noAutofit/></wps:bodyPr></wps:wsp></a:graphicData></a:graphic></wp:anchor></w:drawing></mc:Choice><mc:Fallback><w:pict><v:rect style="position:absolute;width:200pt;height:50pt" o:allowincell="f" stroked="f"><v:textbox><w:txbxContent><w:p><w:r><w:t>Frame content</w:t></w:r></w:p></w:txbxContent></v:textbox><w10:wrap type="square" side="largest"/></v:rect></w:pict></mc:Fallback></mc:AlternateContent></w:r><w:r><w:rPr></w:rPr><w:t>Paragraph with a floating frame.</w:t></w:r></w:p>',
        styles=STYLES_LO.replace("</w:styles>", '<w:style w:type="character" w:styleId="FootnoteAnchor"><w:name w:val="Footnote Anchor"/><w:rPr><w:vertAlign w:val="superscript"/></w:rPr></w:style></w:styles>'), ns=NS_LO, sect=LO_SECT,
        extra={"word/footnotes.xml": ("footnotes", f'{DECL}<w:footnotes {NS_MIN}><w:footnote w:id="0" w:type="separator"><w:p><w:r><w:separator/></w:r></w:p></w:footnote><w:footnote w:id="1" w:type="continuationSeparator"><w:p><w:r><w:continuationSeparator/></w:r></w:p></w:footnote><w:footnote w:id="2"><w:p><w:pPr><w:pStyle w:val="Footnote"/><w:rPr></w:rPr></w:pPr><w:r><w:rPr><w:rStyle w:val="FootnoteCharacters"/></w:rPr><w:footnoteRef/></w:r><w:r><w:rPr></w:rPr><w:t xml:space="preserve"> Writer footnote.</w:t></w:r></w:p></w:footnote></w:footnotes>')},
        rels=[("rId2", "footnotes", "footnotes.xml", False)])

# --- Pathological -----------------------------------------------------------
package("path-pretty-printed.docx",
        p("Pretty-printed XML with newlines between every element.") + p("Second paragraph.", '<w:jc w:val="center"/>'),
        pretty=True, ns=NS_MIN)

package("path-no-styles.docx", '<w:p><w:r><w:t>No styles part, no settings, nothing but the main part.</w:t></w:r></w:p>', styles="", ns=NS_MIN)

package("path-sectpr-only.docx", "", ns=NS_MIN)

package("path-unusual-part-names.docx", p("Main part is not word/document.xml."), main="word/document2.xml", ns=NS_MIN)

package("path-stored-entries.docx", p("All zip entries stored, none deflated."), ns=NS_MIN,
        stored=("[Content_Types].xml", "_rels/.rels", "word/document.xml", "word/styles.xml", "word/_rels/document.xml.rels"))

package("path-bom-and-decl.docx", p("XML declaration variants."), ns=NS_MIN, decl='﻿<?xml version="1.0" encoding="utf-8"?>\n')

package("path-whitespace-text.docx",
        '<w:p><w:r><w:t xml:space="preserve">   leading and trailing   </w:t></w:r><w:r><w:t>no preserve   </w:t></w:r><w:r><w:t xml:space="preserve"></w:t></w:r><w:r><w:t/></w:r><w:r/><w:r><w:rPr><w:b/></w:rPr></w:r><w:r><w:t xml:space="preserve">multi\nline\ttab</w:t></w:r></w:p>'
        + '<w:p><w:r><w:t>&#x2014;&#8212;&amp;&lt;&gt;&quot;&apos;</w:t></w:r><w:r><w:t><![CDATA[cdata <text> & more]]></w:t></w:r></w:p>'
        + '<w:p>\n  <w:r>\n    <w:t>text</w:t>\n  </w:r>\n</w:p>'
        + '<w:p><!-- an XML comment --><w:r><w:t>after comment</w:t></w:r><?pi target?></w:p>',
        ns=NS_MIN)

package("path-namespace-prefixes.docx",
        '<p><r><t>Default namespace, no prefix.</t></r></p><p><pPr><jc val="center"/></pPr><r><rPr><b/></rPr><t>Centered bold.</t></r></p>',
        ns='xmlns="http://schemas.openxmlformats.org/wordprocessingml/2006/main"',
        sect='<sectPr><pgSz w="12240" h="15840"/><pgMar top="1440" right="1440" bottom="1440" left="1440" header="720" footer="720" gutter="0"/></sectPr>',
        styles="")

package("path-many-empty-paras.docx", "".join("<w:p/>" if i % 3 else '<w:p><w:pPr><w:jc w:val="center"/></w:pPr></w:p>' for i in range(300)), ns=NS_MIN)

package("path-deep-nesting.docx",
        '<w:p><w:hyperlink r:id="rId9"><w:ins w:id="1" w:author="A" w:date="2025-01-01T00:00:00Z"><w:sdt><w:sdtPr><w:id w:val="1"/></w:sdtPr><w:sdtContent><w:smartTag w:uri="urn:x" w:element="place"><w:customXml w:uri="urn:y" w:element="z"><w:r><w:rPr><w:b/></w:rPr><w:t>deeply nested run</w:t></w:r></w:customXml></w:smartTag></w:sdtContent></w:sdt></w:ins></w:hyperlink><w:dir w:val="rtl"><w:bdo w:val="ltr"><w:r><w:t>bidi wrappers</w:t></w:r></w:bdo></w:dir></w:p>'
        + '<w:sdt><w:sdtPr><w:id w:val="2"/></w:sdtPr><w:sdtContent><w:sdt><w:sdtPr><w:id w:val="3"/></w:sdtPr><w:sdtContent><w:p><w:r><w:t>nested block controls</w:t></w:r></w:p></w:sdtContent></w:sdt></w:sdtContent></w:sdt>'
        + '<w:customXml w:uri="urn:block" w:element="b"><w:p><w:r><w:t>block-level customXml</w:t></w:r></w:p></w:customXml>'
        + '<w:p><w:r><w:t>Unknown element: </w:t><w:unknownThing w:x="1">inside</w:unknownThing><w:t> after</w:t></w:r><w:unknownParaChild/></w:p><w:unknownBlock><w:p><w:r><w:t>x</w:t></w:r></w:p></w:unknownBlock>',
        ns=NS_MIN, rels=[("rId9", "hyperlink", "https://example.com/", True)])

package("path-huge-paragraph.docx", '<w:p><w:r><w:t xml:space="preserve">' + (LOREM * 400) + '</w:t></w:r></w:p>', ns=NS_MIN)

package("path-many-runs.docx", '<w:p>' + "".join(f'<w:r><w:rPr>{"<w:b/>" if i % 2 else ""}</w:rPr><w:t xml:space="preserve">{c}</w:t></w:r>' for i, c in enumerate(LOREM * 20)) + '</w:p>', ns=NS_MIN)

package("path-duplicate-bookmarks.docx",
        '<w:p><w:bookmarkStart w:id="1" w:name="dup"/><w:r><w:t>one</w:t></w:r><w:bookmarkEnd w:id="1"/></w:p><w:p><w:bookmarkStart w:id="2" w:name="dup"/><w:r><w:t>two</w:t></w:r><w:bookmarkEnd w:id="2"/><w:bookmarkEnd w:id="99"/><w:bookmarkStart w:id="3" w:name="unclosed"/></w:p>',
        ns=NS_MIN)

package("path-missing-rels-targets.docx",
        p("Rels point at parts that do not exist."), ns=NS_MIN,
        rels=[("rId7", "footnotes", "footnotes.xml", False), ("rId8", "header", "header9.xml", False), ("rId9", "numbering", "numbering.xml", False)])

package("path-attr-quotes.docx",
        "<w:p w:rsidR='00A1'><w:pPr><w:jc w:val='right' /></w:pPr><w:r ><w:rPr><w:b  w:val = \"true\"/></w:rPr><w:t xml:space='preserve' >single-quoted attributes and odd spacing</w:t></w:r></w:p>",
        ns=NS_MIN)

package("path-page-breaks-everywhere.docx",
        '<w:p><w:r><w:br w:type="page"/></w:r></w:p>' + p("After leading page break") + '<w:p><w:r><w:br w:type="page"/></w:r><w:r><w:br w:type="page"/></w:r></w:p>' + '<w:p><w:pPr><w:pageBreakBefore/></w:pPr><w:r><w:t>pageBreakBefore</w:t></w:r></w:p>' + '<w:p><w:r><w:t>ends with break</w:t></w:r><w:r><w:br w:type="page"/></w:r></w:p><w:p><w:pPr><w:sectPr><w:type w:val="oddPage"/><w:pgSz w:w="12240" w:h="15840"/><w:pgMar w:top="1440" w:right="1440" w:bottom="1440" w:left="1440" w:header="720" w:footer="720" w:gutter="0"/></w:sectPr></w:pPr></w:p>' + p("last"),
        ns=NS_MIN)

package("path-tab-heavy.docx",
        '<w:p><w:pPr><w:tabs><w:tab w:val="right" w:leader="dot" w:pos="9000"/></w:tabs></w:pPr><w:r><w:t>Chapter 1</w:t><w:tab/><w:t>3</w:t></w:r></w:p>'
        + '<w:p><w:r><w:tab/><w:tab/><w:tab/><w:t>three leading tabs</w:t><w:tab/><w:tab/></w:r></w:p>'
        + '<w:p><w:pPr><w:ind w:left="2880" w:hanging="2880"/></w:pPr><w:r><w:t>Term</w:t><w:tab/><w:t>Definition text that wraps around to the hanging indent position. ' + LOREM + '</w:t></w:r></w:p>',
        ns=NS_MIN)

package("path-mixed-ins-del-para-marks.docx",
        '<w:p><w:pPr><w:rPr><w:ins w:id="1" w:author="A" w:date="2025-01-01T00:00:00Z"/></w:rPr></w:pPr><w:r><w:t>Paragraph mark inserted (paragraphs joined when accepted)</w:t></w:r></w:p><w:p><w:pPr><w:rPr><w:del w:id="2" w:author="A" w:date="2025-01-01T00:00:00Z"/></w:rPr></w:pPr><w:r><w:t>Paragraph mark deleted</w:t></w:r></w:p><w:p><w:r><w:t>Normal</w:t></w:r></w:p>',
        ns=NS_MIN)

package("path-theme-fonts.docx",
        '<w:p><w:r><w:rPr><w:rFonts w:asciiTheme="majorHAnsi" w:hAnsiTheme="majorHAnsi"/></w:rPr><w:t>Major theme font</w:t></w:r><w:r><w:rPr><w:rFonts w:asciiTheme="minorHAnsi"/></w:rPr><w:t xml:space="preserve"> minor theme font</w:t></w:r><w:r><w:rPr><w:rFonts w:cs="Arial" w:eastAsia="MS Mincho"/></w:rPr><w:t xml:space="preserve"> cs/eastAsia only</w:t></w:r><w:r><w:rPr><w:rFonts w:hAnsi="Verdana"/></w:rPr><w:t xml:space="preserve"> hAnsi only</w:t></w:r><w:r><w:rPr><w:sz w:val="1"/></w:rPr><w:t>tiny</w:t></w:r><w:r><w:rPr><w:sz w:val="3276"/></w:rPr><w:t>huge</w:t></w:r><w:r><w:rPr><w:sz w:val="notanumber"/><w:color w:val="zzzzzz"/><w:highlight w:val="notacolor"/><w:u w:val="weird"/></w:rPr><w:t>garbage values</w:t></w:r></w:p>',
        extra={"word/theme/theme1.xml": ("theme", THEME.replace("Calibri Light", "Georgia").replace('typeface="Calibri"', 'typeface="Verdana"'))},
        rels=[("rId3", "theme", "theme/theme1.xml", False)])

package("path-extra-zip-entries.docx",
        p("Zip has extra unrelated entries and a directory entry."), ns=NS_MIN)
with zipfile.ZipFile(f"{out_dir}/path-extra-zip-entries.docx") as z:
    parts = {n: z.read(n) for n in z.namelist()}
parts["word/media/"] = b""
parts["word/media/blob.bin"] = bytes(range(256)) * 4
parts["customXml/itemProps1.xml"] = b'<?xml version="1.0"?><ds:datastoreItem xmlns:ds="http://schemas.openxmlformats.org/officeDocument/2006/customXml"/>'
parts["docProps/thumbnail.jpeg"] = b"\xff\xd8\xff\xe0" + b"\x00" * 64 + b"\xff\xd9"
save("path-extra-zip-entries.docx", parts, stored=("word/media/blob.bin",))


# --- The original 0.1 hand-built fixture: a bit of everything in one file ---
W = 'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml" xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006" mc:Ignorable="w14"'
DOC = f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document {W}><w:body>
<w:p w:rsidR="00A1" w14:paraId="1"><w:pPr><w:pStyle w:val="Heading1"/><w:rPr><w:lang w:val="en-GB"/></w:rPr></w:pPr><w:r><w:rPr><w:lang w:val="en-GB"/></w:rPr><w:t>Pathological</w:t></w:r></w:p>
<w:p><w:r><w:t xml:space="preserve">Plain </w:t></w:r><w:r><w:rPr><w:b/><w:bCs/><w:rFonts w:ascii="Arial" w:hAnsi="Arial" w:cs="Arial"/></w:rPr><w:t>bold Arial</w:t></w:r><w:r><w:t xml:space="preserve"> then </w:t></w:r><w:ins w:id="1" w:author="A" w:date="2024-01-01T00:00:00Z"><w:r><w:t>inserted</w:t></w:r></w:ins><w:del w:id="2" w:author="A" w:date="2024-01-01T00:00:00Z"><w:r><w:delText xml:space="preserve"> deleted</w:delText></w:r></w:del><w:r><w:t>.</w:t></w:r></w:p>
<w:p><w:commentRangeStart w:id="0"/><w:r><w:t>Commented text</w:t></w:r><w:commentRangeEnd w:id="0"/><w:r><w:commentReference w:id="0"/></w:r></w:p>
<w:p><w:r><w:t xml:space="preserve">See </w:t></w:r><w:hyperlink r:id="rId9" w:history="1"><w:r><w:rPr><w:rStyle w:val="Hyperlink"/></w:rPr><w:t>the website</w:t></w:r></w:hyperlink><w:r><w:t xml:space="preserve"> and page </w:t></w:r><w:r><w:fldChar w:fldCharType="begin"/></w:r><w:r><w:instrText xml:space="preserve"> PAGE </w:instrText></w:r><w:r><w:fldChar w:fldCharType="separate"/></w:r><w:r><w:t>3</w:t></w:r><w:r><w:fldChar w:fldCharType="end"/></w:r><w:r><w:t>.</w:t></w:r></w:p>
<w:bookmarkStart w:id="7" w:name="regional"/><w:p><w:pPr><w:jc w:val="center"/><w:ind w:left="720" w:hanging="360"/><w:spacing w:before="240" w:after="0" w:line="360" w:lineRule="auto"/><w:tabs><w:tab w:val="right" w:leader="dot" w:pos="8640"/></w:tabs></w:pPr><w:r><w:rPr><w:i/><w:sz w:val="28"/><w:color w:val="FF0000"/><w:highlight w:val="yellow"/><w:u w:val="double"/></w:rPr><w:t>Formatted</w:t></w:r><w:r><w:tab/><w:t>tabbed</w:t><w:br/><w:t>after break</w:t></w:r></w:p><w:bookmarkEnd w:id="7"/>
<w:sdt><w:sdtPr><w:alias w:val="Title"/><w:id w:val="123"/></w:sdtPr><w:sdtContent><w:p><w:r><w:t>Inside a block content control</w:t></w:r></w:p></w:sdtContent></w:sdt>
<w:p><w:r><w:t xml:space="preserve">Run-level </w:t></w:r><w:sdt><w:sdtPr><w:id w:val="124"/></w:sdtPr><w:sdtContent><w:r><w:t>control</w:t></w:r></w:sdtContent></w:sdt><w:r><w:t xml:space="preserve"> here, then a footnote</w:t></w:r><w:r><w:rPr><w:vertAlign w:val="superscript"/></w:rPr><w:footnoteReference w:id="1"/></w:r><w:r><w:t>.</w:t></w:r></w:p>
<w:tbl><w:tblPr><w:tblStyle w:val="TableGrid"/><w:tblW w:w="0" w:type="auto"/></w:tblPr><w:tblGrid><w:gridCol w:w="4000"/><w:gridCol w:w="4000"/></w:tblGrid><w:tr><w:tc><w:tcPr><w:tcW w:w="4000" w:type="dxa"/></w:tcPr><w:p><w:r><w:t>A1</w:t></w:r></w:p></w:tc><w:tc><w:tcPr><w:tcW w:w="4000" w:type="dxa"/></w:tcPr><w:p><w:r><w:t>B1</w:t></w:r></w:p></w:tc></w:tr></w:tbl>
<w:p><w:r><w:t>Last &amp; &lt;escaped&gt; paragraph</w:t></w:r><w:r><w:br w:type="page"/></w:r><w:r><w:t>New page</w:t></w:r></w:p>
<w:sectPr w:rsidR="00B2"><w:footerReference w:type="default" r:id="rId8"/><w:pgSz w:w="11906" w:h="16838"/><w:pgMar w:top="1440" w:right="1800" w:bottom="1440" w:left="1800" w:header="708" w:footer="708" w:gutter="0"/><w:cols w:space="708"/><w:docGrid w:linePitch="360"/></w:sectPr>
</w:body></w:document>'''

STYLES = f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:styles {W}><w:docDefaults><w:rPrDefault><w:rPr><w:rFonts w:asciiTheme="minorHAnsi" w:hAnsiTheme="minorHAnsi"/><w:sz w:val="24"/><w:lang w:val="en-US"/></w:rPr></w:rPrDefault><w:pPrDefault><w:pPr><w:spacing w:after="200" w:line="276" w:lineRule="auto"/></w:pPr></w:pPrDefault></w:docDefaults>
<w:latentStyles w:defLockedState="0" w:count="1"><w:lsdException w:name="Normal" w:uiPriority="0"/></w:latentStyles>
<w:style w:type="paragraph" w:default="1" w:styleId="Normal"><w:name w:val="Normal"/><w:qFormat/></w:style>
<w:style w:type="paragraph" w:styleId="Heading1"><w:name w:val="heading 1"/><w:basedOn w:val="Normal"/><w:next w:val="Normal"/><w:pPr><w:keepNext/><w:spacing w:before="480" w:after="0"/><w:outlineLvl w:val="0"/></w:pPr><w:rPr><w:rFonts w:asciiTheme="majorHAnsi" w:hAnsiTheme="majorHAnsi"/><w:b/><w:color w:val="365F91"/><w:sz w:val="28"/></w:rPr></w:style>
<w:style w:type="character" w:default="1" w:styleId="DefaultParagraphFont"><w:name w:val="Default Paragraph Font"/><w:semiHidden/></w:style>
<w:style w:type="character" w:styleId="Hyperlink"><w:name w:val="Hyperlink"/><w:rPr><w:color w:val="0000FF"/><w:u w:val="single"/></w:rPr></w:style>
</w:styles>'''

FOOTNOTES = f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:footnotes {W}><w:footnote w:type="separator" w:id="-1"><w:p><w:r><w:separator/></w:r></w:p></w:footnote><w:footnote w:id="1"><w:p><w:pPr><w:pStyle w:val="FootnoteText"/></w:pPr><w:r><w:footnoteRef/></w:r><w:r><w:t xml:space="preserve"> A footnote.</w:t></w:r></w:p></w:footnote></w:footnotes>'''

COMMENTS = f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:comments {W}><w:comment w:id="0" w:author="Reviewer" w:date="2024-01-01T00:00:00Z" w:initials="R"><w:p><w:r><w:t>Please check this.</w:t></w:r></w:p></w:comment></w:comments>'''

FOOTER = f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:ftr {W}><w:p><w:pPr><w:jc w:val="center"/></w:pPr><w:r><w:fldChar w:fldCharType="begin"/></w:r><w:r><w:instrText>PAGE</w:instrText></w:r><w:r><w:fldChar w:fldCharType="end"/></w:r></w:p></w:ftr>'''

THEME = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<a:theme xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" name="Office Theme"><a:themeElements><a:fontScheme name="Office"><a:majorFont><a:latin typeface="Cambria"/></a:majorFont><a:minorFont><a:latin typeface="Times New Roman"/></a:minorFont></a:fontScheme></a:themeElements></a:theme>'''

CT = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types"><Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/><Default Extension="xml" ContentType="application/xml"/><Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/><Override PartName="/word/styles.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.styles+xml"/><Override PartName="/word/footnotes.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml"/><Override PartName="/word/comments.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"/><Override PartName="/word/footer1.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.footer+xml"/><Override PartName="/word/theme/theme1.xml" ContentType="application/vnd.openxmlformats-officedocument.theme+xml"/></Types>'''

RELS = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"><Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/></Relationships>'''

DOC_RELS = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"><Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/styles" Target="styles.xml"/><Relationship Id="rId2" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes" Target="footnotes.xml"/><Relationship Id="rId3" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments" Target="comments.xml"/><Relationship Id="rId4" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/theme" Target="theme/theme1.xml"/><Relationship Id="rId8" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/footer" Target="footer1.xml"/><Relationship Id="rId9" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink" Target="https://example.com/" TargetMode="External"/></Relationships>'''

with zipfile.ZipFile(f"{out_dir}/path-mixed.docx", "w", zipfile.ZIP_DEFLATED) as z:
    z.writestr("[Content_Types].xml", CT)
    z.writestr("_rels/.rels", RELS)
    z.writestr("word/document.xml", DOC)
    z.writestr("word/_rels/document.xml.rels", DOC_RELS)
    z.writestr("word/styles.xml", STYLES)
    z.writestr("word/footnotes.xml", FOOTNOTES)
    z.writestr("word/comments.xml", COMMENTS)
    z.writestr("word/footer1.xml", FOOTER)
    z.writestr("word/theme/theme1.xml", THEME)
    z.writestr("word/media/blob.bin", b"\x00\x01\x02binary")
written.append("path-mixed.docx")

with open(f"{out_dir}/README.md", "w") as f:
    f.write("# Round-trip corpus\n\nGenerated by `tools/make_corpus.py`; every file here must pass\n`crates/wp-docx/tests/roundtrip.rs` (SPEC §10.2). Do not edit by hand —\nchange the generator and regenerate.\n\n"
            "| Prefix | Origin |\n|---|---|\n| `gen-` | python-docx |\n| `word-` | hand-built to match Microsoft Word 2016–365 output |\n| `gdocs-` | hand-built to match Google Docs export |\n| `lo-` | hand-built to match LibreOffice Writer export |\n| `path-` | deliberately pathological |\n\n"
            + "\n".join(f"- `{n}`" for n in sorted(written)) + "\n")
print(f"wrote {len(written)} files")
